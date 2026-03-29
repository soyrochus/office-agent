from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from offagent.domain.locators import parse_locator, to_v2_locator
from offagent.domain.models import (
    BlockStyle,
    BlockBundle,
    DocxRun,
    DocxParagraph,
    DocxTableCell,
    DocxTable,
    DocumentBlock,
    DocumentRef,
    InlineFragment,
    InlineStyle,
    IndexedItem,
    SectionPayload,
    StructureSection,
    TextContainerSnapshot,
    VisibleTextRange,
)
from offagent.domain.text_fragments import (
    apply_style_to_range,
    fragment_text,
    normalize_fragments,
)
from offagent.errors import InvalidArgumentsError, TargetNotEditableError, TargetNotFoundError

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Pt, RGBColor
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ModuleNotFoundError:  # pragma: no cover - exercised through dependency checks
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_COLOR_INDEX = None
    CT_Tbl = None
    CT_P = None
    Pt = None
    RGBColor = None
    Table = None
    Paragraph = None
    Run = None

RunFormatting = InlineStyle


@dataclass(frozen=True)
class ResolvedParagraphTarget:
    block_index: int
    paragraph_index: int
    paragraph: Paragraph


@dataclass(frozen=True)
class ResolvedTableCellTarget:
    block_index: int
    table_index: int
    row_index: int
    column_index: int
    table: Table


ResolvedTarget = ResolvedParagraphTarget | ResolvedTableCellTarget


def extract_document(document_path: Path) -> list[IndexedItem]:
    items: list[IndexedItem] = []

    for paragraph in get_paragraphs(document_path):
        locator = f"para:{paragraph.paragraph_index}"
        items.append(
            IndexedItem(
                item_id=locator,
                item_type="paragraph",
                locator=locator,
                preview=paragraph.preview,
                content_text=paragraph.text,
                metadata={
                    "paragraph_index": paragraph.paragraph_index,
                    "block_index": paragraph.block_index,
                    "style_name": paragraph.style_name,
                    "is_heading": paragraph.is_heading,
                },
            )
        )

    return items


def build_embedding_text(item: IndexedItem, document_path: Path) -> str:
    del document_path
    return item.content_text


def read_paragraph(document_path: Path, item_id: str) -> str:
    paragraph = _resolve_paragraph(_open_document(document_path), item_id)
    return paragraph.text


def replace_paragraph(document_path: Path, item_id: str, text: str, output_path: Path | None = None) -> Path:
    document = _open_document(document_path)
    paragraph = _resolve_paragraph(document, item_id)
    formatting = _capture_run_formatting(paragraph.runs[0] if paragraph.runs else None)
    _clear_paragraph(paragraph)
    replacement_run = paragraph.add_run(text)
    _apply_run_formatting(replacement_run, formatting)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path


def append_paragraph(document_path: Path, item_id: str, text: str, output_path: Path | None = None) -> Path:
    document = _open_document(document_path)
    paragraph = _resolve_paragraph(document, item_id)
    if paragraph.runs:
        paragraph.runs[-1].text = f"{paragraph.runs[-1].text}{text}"
    else:
        paragraph.add_run(text)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path


def make_table_cell_locator(table_index: int, row_index: int, column_index: int) -> str:
    return f"table:{table_index}:cell:{row_index}:{column_index}"


def parse_table_cell_locator(locator: str) -> tuple[int, int, int]:
    parts = locator.split(":")
    if len(parts) != 5 or parts[0] != "table" or parts[2] != "cell":
        raise InvalidArgumentsError(f"Unsupported DOCX table cell locator: {locator}")
    try:
        table_index = int(parts[1])
        row_index = int(parts[3])
        column_index = int(parts[4])
    except ValueError as exc:
        raise InvalidArgumentsError(f"Invalid DOCX table cell locator: {locator}") from exc
    return table_index, row_index, column_index


def resolve_structure(document_path: Path) -> tuple[StructureSection, ...]:
    document = _open_document(document_path)
    sections: list[StructureSection] = []

    paragraph_index = 0
    table_index = 0
    for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if block_type == "paragraph":
            paragraph_model = _paragraph_model(block, block_index, paragraph_index)
            sections.append(
                StructureSection(
                    locator=f"para:{paragraph_index}",
                    section_type="paragraph",
                    preview=paragraph_model.preview,
                    metadata={
                        "block_index": block_index,
                        "block_type": "paragraph",
                        "paragraph_index": paragraph_index,
                        "style_name": paragraph_model.style_name,
                        "is_heading": paragraph_model.is_heading,
                    },
                )
            )
            paragraph_index += 1
            continue

        table_model = _table_model(block, block_index, table_index)
        sections.append(
            StructureSection(
                locator=make_table_cell_locator(table_index, 0, 0),
                section_type="table",
                preview=table_model.preview,
                metadata={
                    "block_index": block_index,
                    "block_type": "table",
                    "table_index": table_index,
                    "row_count": len(table_model.rows),
                    "column_count": max((len(row) for row in table_model.rows), default=0),
                },
            )
        )
        table_index += 1

    return tuple(sections)


def get_section(document_path: Path, locator: str) -> SectionPayload:
    document = _open_document(document_path)
    resolved = _resolve_locator(document, locator)
    document_ref = _document_ref(document_path)

    if isinstance(resolved, ResolvedParagraphTarget):
        paragraph_model = _paragraph_model(
            resolved.paragraph,
            resolved.block_index,
            resolved.paragraph_index,
        )
        return SectionPayload(
            document=document_ref,
            locator=f"para:{resolved.paragraph_index}",
            section_type="paragraph",
            preview=paragraph_model.preview,
            metadata={
                "block_index": resolved.block_index,
                "block_type": "paragraph",
                "paragraph_index": resolved.paragraph_index,
            },
            block_type="paragraph",
            text=paragraph_model.text,
            style_name=paragraph_model.style_name,
            is_heading=paragraph_model.is_heading,
            runs=tuple(_run_model(run) for run in resolved.paragraph.runs),
        )

    table_model = _table_model(resolved.table, resolved.block_index, resolved.table_index)
    cells = tuple(
        DocxTableCell(
            locator=make_table_cell_locator(resolved.table_index, row_index, column_index),
            row_index=row_index,
            column_index=column_index,
            text=cell.text,
            metadata={},
        )
        for row_index, row in enumerate(resolved.table.rows)
        for column_index, cell in enumerate(row.cells)
    )
    return SectionPayload(
        document=document_ref,
        locator=make_table_cell_locator(resolved.table_index, 0, 0),
        section_type="table",
        preview=table_model.preview,
        metadata={
            "block_index": resolved.block_index,
            "block_type": "table",
            "table_index": resolved.table_index,
            "row_count": len(table_model.rows),
            "column_count": max((len(row) for row in table_model.rows), default=0),
        },
        block_type="table",
        rows=table_model.rows,
        table_cells=cells,
    )


def read_node(document_path: Path, locator: str) -> tuple[str, str, dict[str, object]]:
    document = _open_document(document_path)
    resolved = _resolve_locator(document, locator)

    if isinstance(resolved, ResolvedParagraphTarget):
        paragraph_model = _paragraph_model(
            resolved.paragraph,
            resolved.block_index,
            resolved.paragraph_index,
        )
        return (
            "paragraph",
            paragraph_model.text,
            {
                "block_index": resolved.block_index,
                "paragraph_index": resolved.paragraph_index,
                "style_name": paragraph_model.style_name,
                "is_heading": paragraph_model.is_heading,
            },
        )

    cell = resolved.table.rows[resolved.row_index].cells[resolved.column_index]
    return (
        "table_cell",
        cell.text,
        {
            "block_index": resolved.block_index,
            "table_index": resolved.table_index,
            "row_index": resolved.row_index,
            "column_index": resolved.column_index,
        },
    )


def write_node(document_path: Path, locator: str, text: str, output_path: Path | None = None) -> Path:
    document = _open_document(document_path)
    resolved = _resolve_locator(document, locator)

    if isinstance(resolved, ResolvedParagraphTarget):
        return replace_paragraph(document_path, f"para:{resolved.paragraph_index}", text, output_path)

    cell = resolved.table.rows[resolved.row_index].cells[resolved.column_index]
    cell.text = text
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path


def read_paragraph_fragments(document_path: Path, locator: str) -> TextContainerSnapshot:
    document = _open_document(document_path)
    canonical, components = _canonical_docx_locator(locator)
    if len(components) != 3 or components[:2] != ("docx", "para"):
        raise InvalidArgumentsError("DOCX fragment reads require a paragraph locator.")

    paragraph = _resolve_paragraph(document, f"para:{components[2]}")
    fragments = _read_docx_paragraph_fragments(paragraph)
    return TextContainerSnapshot(
        locator=canonical,
        object_type="paragraph",
        text=fragment_text(fragments),
        fragments=fragments,
        metadata={"paragraph_index": int(components[2])},
    )


def rewrite_paragraph_fragments(
    document_path: Path,
    locator: str,
    fragments: list[InlineFragment] | tuple[InlineFragment, ...],
    output_path: Path | None = None,
) -> tuple[Path, str, TextContainerSnapshot]:
    document = _open_document(document_path)
    canonical, components = _canonical_docx_locator(locator)
    if len(components) != 3 or components[:2] != ("docx", "para"):
        raise InvalidArgumentsError("DOCX fragment writes require a paragraph locator.")

    paragraph = _resolve_paragraph(document, f"para:{components[2]}")
    normalized = normalize_fragments(fragments)
    _rewrite_docx_paragraph(paragraph, normalized)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    snapshot = TextContainerSnapshot(
        locator=canonical,
        object_type="paragraph",
        text=fragment_text(normalized),
        fragments=normalized,
        metadata={"paragraph_index": int(components[2])},
    )
    return target_path, canonical, snapshot


def insert_paragraph(
    document_path: Path,
    text: str,
    *,
    style_name: str | None = None,
    after_locator: str | None = None,
    output_path: Path | None = None,
) -> tuple[Path, str]:
    if after_locator is None:
        target_path, block_index = append_paragraph_block(
            document_path,
            text,
            style_name=style_name,
            output_path=output_path,
        )
        paragraph_count = len(get_paragraphs(target_path))
        return target_path, f"para:{paragraph_count - 1}"

    document = _open_document(document_path)
    resolved = _resolve_locator(document, after_locator)
    paragraph_index = (
        resolved.paragraph_index + 1
        if isinstance(resolved, ResolvedParagraphTarget)
        else _paragraphs_before_block(document, resolved.block_index)
    )

    anchor_element = (
        resolved.paragraph._element
        if isinstance(resolved, ResolvedParagraphTarget)
        else resolved.table._element
    )
    new_element = document.element.body.add_p()
    anchor_element.addnext(new_element)
    paragraph = Paragraph(new_element, document)
    paragraph.add_run(text)
    if style_name is not None:
        try:
            paragraph.style = style_name
        except (KeyError, ValueError) as exc:
            raise InvalidArgumentsError(f"Unknown DOCX paragraph style: {style_name}") from exc

    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, f"para:{paragraph_index}"


def create_docx(output_path: Path) -> Path:
    document = _open_document_from_default_template()
    document.save(output_path)
    return output_path


def add_paragraph(
    document_path: Path,
    text: str,
    output_path: Path | None = None,
) -> tuple[Path, str]:
    target_path, legacy_locator = insert_paragraph(
        document_path,
        text,
        output_path=output_path,
    )
    return target_path, to_v2_locator(legacy_locator, file_type="docx")


def add_heading(
    document_path: Path,
    text: str,
    level: int,
    output_path: Path | None = None,
) -> tuple[Path, str]:
    if level < 1 or level > 9:
        raise InvalidArgumentsError("DOCX heading level must be between 1 and 9.")

    document = _open_document(document_path)
    paragraph = document.add_heading(text, level=level)
    paragraph_index = sum(1 for block_type, _ in _iter_blocks(document) if block_type == "paragraph") - 1
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, f"docx:para:{paragraph_index}"


def add_table(
    document_path: Path,
    rows: int,
    columns: int,
    output_path: Path | None = None,
) -> tuple[Path, str]:
    if rows < 1 or columns < 1:
        raise InvalidArgumentsError("DOCX table rows and columns must be positive.")

    document = _open_document(document_path)
    table_index = sum(1 for block_type, _ in _iter_blocks(document) if block_type == "table")
    document.add_table(rows=rows, cols=columns)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, f"docx:table:{table_index}"


def style_run(
    document_path: Path,
    locator: str,
    style: InlineStyle,
    clear_fields: list[str] | tuple[str, ...],
    output_path: Path | None = None,
) -> tuple[Path, str, dict[str, object]]:
    document = _open_document(document_path)
    canonical, components = _canonical_docx_locator(locator)
    if len(components) != 5 or components[:2] != ("docx", "para") or components[3] != "run":
        raise InvalidArgumentsError("DOCX inline styling requires a run locator.")

    paragraph = _resolve_paragraph(document, f"para:{components[2]}")
    run_index = _parse_int_component(components[4], locator)
    try:
        run = paragraph.runs[run_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Run {run_index} does not exist in paragraph {components[2]}."
        ) from exc

    cleared_fields = _normalize_clear_fields(clear_fields, _INLINE_STYLE_FIELDS)
    _apply_docx_inline_style(run, style, cleared_fields)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, canonical, {"cleared_fields": cleared_fields}


def style_paragraph_range(
    document_path: Path,
    locator: str,
    text_range: VisibleTextRange,
    style: InlineStyle,
    clear_fields: list[str] | tuple[str, ...],
    output_path: Path | None = None,
) -> tuple[Path, str, dict[str, object]]:
    snapshot = read_paragraph_fragments(document_path, locator)
    cleared_fields = _normalize_clear_fields(clear_fields, _INLINE_STYLE_FIELDS)
    styled = apply_style_to_range(snapshot.fragments, text_range, style=style, clear_fields=cleared_fields)
    target_path, canonical, rewritten = rewrite_paragraph_fragments(
        document_path,
        locator,
        styled,
        output_path=output_path,
    )
    return target_path, canonical, {
        "cleared_fields": cleared_fields,
        "range": {"start": text_range.start, "end": text_range.end},
        "text": rewritten.text,
    }


def style_paragraph(
    document_path: Path,
    locator: str,
    style: BlockStyle,
    clear_fields: list[str] | tuple[str, ...],
    output_path: Path | None = None,
) -> tuple[Path, str, dict[str, object]]:
    document = _open_document(document_path)
    canonical, components = _canonical_docx_locator(locator)
    if len(components) != 3 or components[:2] != ("docx", "para"):
        raise InvalidArgumentsError("DOCX block styling requires a paragraph locator.")

    paragraph = _resolve_paragraph(document, f"para:{components[2]}")
    cleared_fields = _normalize_clear_fields(clear_fields, _BLOCK_STYLE_FIELDS)
    skipped_fields = _apply_docx_block_style(paragraph, style, cleared_fields)
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, canonical, {"cleared_fields": cleared_fields, "skipped_fields": skipped_fields}


def set_structural_role(
    document_path: Path,
    locator: str,
    role: str,
    level: int | None,
    output_path: Path | None = None,
) -> tuple[Path, str, dict[str, object]]:
    document = _open_document(document_path)
    canonical, components = _canonical_docx_locator(locator)
    if len(components) != 3 or components[:2] != ("docx", "para"):
        raise InvalidArgumentsError("set_structural_role requires a DOCX paragraph locator.")

    style_name = _docx_structural_style_name(role, level)
    if not any(getattr(style, "name", None) == style_name for style in document.styles):
        raise TargetNotEditableError(f"DOCX style {style_name!r} is not available in the document.")

    paragraph = _resolve_paragraph(document, f"para:{components[2]}")
    paragraph.style = style_name
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, canonical, {"role": role, "level": level, "style_name": style_name}


def get_blocks(document_path: Path) -> tuple[DocumentBlock, ...]:
    document = _open_document(document_path)
    blocks: list[DocumentBlock] = []

    paragraph_index = 0
    table_index = 0
    for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if block_type == "paragraph":
            paragraph_model = _paragraph_model(block, block_index, paragraph_index)
            blocks.append(
                DocumentBlock(
                    block_index=block_index,
                    block_type="paragraph",
                    preview=paragraph_model.preview,
                    metadata={
                        "paragraph_index": paragraph_model.paragraph_index,
                        "style_name": paragraph_model.style_name,
                        "is_heading": paragraph_model.is_heading,
                    },
                )
            )
            paragraph_index += 1
        else:
            table_model = _table_model(block, block_index, table_index)
            blocks.append(
                DocumentBlock(
                    block_index=block_index,
                    block_type="table",
                    preview=table_model.preview,
                    metadata={
                        "table_index": table_model.table_index,
                        "row_count": len(table_model.rows),
                        "column_count": max((len(row) for row in table_model.rows), default=0),
                    },
                )
            )
            table_index += 1

    return tuple(blocks)


def get_paragraphs(document_path: Path) -> tuple[DocxParagraph, ...]:
    document = _open_document(document_path)
    paragraphs: list[DocxParagraph] = []

    paragraph_index = 0
    for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if block_type != "paragraph":
            continue
        paragraphs.append(_paragraph_model(block, block_index, paragraph_index))
        paragraph_index += 1

    return tuple(paragraphs)


def get_tables(document_path: Path) -> tuple[DocxTable, ...]:
    document = _open_document(document_path)
    tables: list[DocxTable] = []

    table_index = 0
    for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if block_type != "table":
            continue
        tables.append(_table_model(block, block_index, table_index))
        table_index += 1

    return tuple(tables)


def get_block_bundle(document_path: Path, block_index: int) -> BlockBundle:
    document = _open_document(document_path)

    paragraph_index = 0
    table_index = 0
    for current_block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if current_block_index != block_index:
            if block_type == "paragraph":
                paragraph_index += 1
            else:
                table_index += 1
            continue

        if block_type == "paragraph":
            paragraph_model = _paragraph_model(block, current_block_index, paragraph_index)
            return BlockBundle(
                document=_document_ref(document_path),
                block=DocumentBlock(
                    block_index=current_block_index,
                    block_type="paragraph",
                    preview=paragraph_model.preview,
                    metadata={
                        "paragraph_index": paragraph_model.paragraph_index,
                        "style_name": paragraph_model.style_name,
                        "is_heading": paragraph_model.is_heading,
                    },
                ),
                paragraph=paragraph_model,
            )

        table_model = _table_model(block, current_block_index, table_index)
        return BlockBundle(
            document=_document_ref(document_path),
            block=DocumentBlock(
                block_index=current_block_index,
                block_type="table",
                preview=table_model.preview,
                metadata={
                    "table_index": table_model.table_index,
                    "row_count": len(table_model.rows),
                    "column_count": max((len(row) for row in table_model.rows), default=0),
                },
            ),
            table=table_model,
        )

    raise TargetNotFoundError(f"Block {block_index} does not exist in the document.")


def append_paragraph_block(
    document_path: Path,
    text: str,
    *,
    style_name: str | None = None,
    output_path: Path | None = None,
) -> tuple[Path, int]:
    document = _open_document(document_path)
    block_index = len(list(_iter_blocks(document)))
    paragraph = document.add_paragraph(text)
    if style_name is not None:
        try:
            paragraph.style = style_name
        except (KeyError, ValueError) as exc:
            raise InvalidArgumentsError(f"Unknown DOCX paragraph style: {style_name}") from exc
    target_path = _target_path(document_path, output_path)
    document.save(target_path)
    return target_path, block_index


def replace_block(document_path: Path, block_index: int, text: str, output_path: Path | None = None) -> Path:
    document = _open_document(document_path)

    paragraph_index = 0
    for current_block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if current_block_index != block_index:
            if block_type == "paragraph":
                paragraph_index += 1
            continue

        if block_type == "table":
            raise TargetNotEditableError("DOCX table block replacement is not supported.")

        item_id = f"para:{paragraph_index}"
        return replace_paragraph(document_path, item_id, text, output_path)

    raise TargetNotFoundError(f"Block {block_index} does not exist in the document.")


def get_tables_result(document_path: Path) -> tuple[DocxTable, ...]:
    return get_tables(document_path)


def _open_document(document_path: Path):
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX operations.")
    return Document(str(document_path))


def _open_document_from_default_template():
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX operations.")
    return Document()


def _document_ref(document_path: Path):
    resolved_path = document_path.resolve()
    stat = resolved_path.stat()
    return DocumentRef(
        document_id=resolved_path.as_posix(),
        path=resolved_path,
        file_type="docx",
        display_name=resolved_path.name,
        modified_time=stat.st_mtime,
    )


def _resolve_paragraph(document, item_id: str):
    if not item_id.startswith("para:"):
        raise InvalidArgumentsError(f"Unsupported DOCX paragraph item id: {item_id}")

    try:
        paragraph_index = int(item_id.split(":", maxsplit=1)[1])
    except ValueError as exc:
        raise InvalidArgumentsError(f"Invalid DOCX paragraph item id: {item_id}") from exc

    try:
        return document.paragraphs[paragraph_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Paragraph {paragraph_index} does not exist in the document."
        ) from exc


def _resolve_locator(document, locator: str) -> ResolvedTarget:
    normalized = locator.strip()
    if normalized.startswith("para:"):
        paragraph_index = _parse_paragraph_locator(normalized)
        current_paragraph_index = 0
        for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
            if block_type != "paragraph":
                continue
            if current_paragraph_index == paragraph_index:
                return ResolvedParagraphTarget(
                    block_index=block_index,
                    paragraph_index=paragraph_index,
                    paragraph=block,
                )
            current_paragraph_index += 1
        raise TargetNotFoundError(f"Paragraph {paragraph_index} does not exist in the document.")

    table_index, row_index, column_index = parse_table_cell_locator(normalized)
    current_table_index = 0
    for block_index, (block_type, block) in enumerate(_iter_blocks(document)):
        if block_type != "table":
            continue
        if current_table_index == table_index:
            try:
                block.rows[row_index].cells[column_index]
            except IndexError as exc:
                raise TargetNotFoundError(
                    f"Table cell {table_index}:{row_index}:{column_index} does not exist."
                ) from exc
            return ResolvedTableCellTarget(
                block_index=block_index,
                table_index=table_index,
                row_index=row_index,
                column_index=column_index,
                table=block,
            )
        current_table_index += 1

    raise TargetNotFoundError(f"Table {table_index} does not exist in the document.")


def _parse_paragraph_locator(locator: str) -> int:
    try:
        return int(locator.split(":", maxsplit=1)[1])
    except ValueError as exc:
        raise InvalidArgumentsError(f"Invalid DOCX paragraph item id: {locator}") from exc


def _capture_run_formatting(run: Run | None) -> RunFormatting | None:
    if run is None:
        return None

    return RunFormatting(
        bold=run.bold,
        italic=run.italic,
        underline=run.underline,
        strike=run.font.strike,
        font_name=run.font.name,
        font_size=None if run.font.size is None else run.font.size.pt,
        font_color=None if run.font.color.rgb is None else str(run.font.color.rgb),
        highlight=_docx_highlight_name(run.font.highlight_color),
    )


def _apply_run_formatting(run: Run, formatting: RunFormatting | None) -> None:
    if formatting is None:
        return

    _apply_docx_inline_style(run, formatting, ())


def _clear_paragraph(paragraph) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag.endswith("}pPr"):
            continue
        paragraph_element.remove(child)


def _ensure_rewritable_docx_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag.endswith("}pPr"):
            continue
        if not child.tag.endswith("}r"):
            raise TargetNotEditableError(
                "DOCX paragraph contains inline content that cannot be safely reconstructed."
            )


def _read_docx_paragraph_fragments(paragraph) -> tuple[InlineFragment, ...]:
    _ensure_rewritable_docx_paragraph(paragraph)
    if not paragraph.runs:
        return ()
    return normalize_fragments(
        [
            InlineFragment(
                text=run.text,
                style=_capture_run_formatting(run) or InlineStyle(),
            )
            for run in paragraph.runs
        ]
    )


def _rewrite_docx_paragraph(
    paragraph,
    fragments: list[InlineFragment] | tuple[InlineFragment, ...],
) -> None:
    _ensure_rewritable_docx_paragraph(paragraph)
    _clear_paragraph(paragraph)
    normalized = normalize_fragments(fragments)
    if not normalized:
        paragraph.add_run("")
        return
    for fragment in normalized:
        run = paragraph.add_run(fragment.text)
        _apply_docx_inline_style(run, fragment.style, ())


def _target_path(document_path: Path, output_path: Path | None) -> Path:
    return document_path if output_path is None else output_path


def _iter_blocks(document) -> list[tuple[str, Paragraph | Table]]:
    if CT_P is None or CT_Tbl is None or Paragraph is None or Table is None:
        raise RuntimeError("python-docx is required for DOCX operations.")

    parent = document.element.body
    blocks: list[tuple[str, Paragraph | Table]] = []
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("paragraph", Paragraph(child, document)))
        elif isinstance(child, CT_Tbl):
            blocks.append(("table", Table(child, document)))
    return blocks


def _paragraph_model(paragraph, block_index: int, paragraph_index: int) -> DocxParagraph:
    style_name = paragraph.style.name if paragraph.style is not None else None
    is_heading = bool(style_name and style_name.startswith("Heading"))
    text = paragraph.text
    return DocxParagraph(
        block_index=block_index,
        paragraph_index=paragraph_index,
        text=text,
        style_name=style_name,
        is_heading=is_heading,
        preview=text[:120],
        metadata={},
    )


def _table_model(table, block_index: int, table_index: int) -> DocxTable:
    rows = tuple(tuple(cell.text for cell in row.cells) for row in table.rows)
    preview = " | ".join(cell for row in rows for cell in row if cell)[:120]
    return DocxTable(
        block_index=block_index,
        table_index=table_index,
        rows=rows,
        preview=preview,
        metadata={},
    )


def _run_model(run) -> DocxRun:
    color_rgb = None
    if run.font.color.rgb is not None:
        color_rgb = str(run.font.color.rgb)

    font_size = None
    if run.font.size is not None:
        font_size = int(run.font.size)

    return DocxRun(
        text=run.text,
        bold=run.bold,
        italic=run.italic,
        underline=run.underline,
        strike=run.font.strike,
        font_name=run.font.name,
        font_size=font_size,
        color_rgb=color_rgb,
    )


def _paragraphs_before_block(document, block_index: int) -> int:
    count = 0
    for current_block_index, (block_type, _) in enumerate(_iter_blocks(document)):
        if current_block_index > block_index:
            break
        if current_block_index == block_index:
            break
        if block_type == "paragraph":
            count += 1
    return count


_INLINE_STYLE_FIELDS = frozenset(
    {
        "bold",
        "italic",
        "underline",
        "strike",
        "font_name",
        "font_size",
        "font_color",
        "highlight",
    }
)
_BLOCK_STYLE_FIELDS = frozenset(
    {
        "alignment",
        "indent_level",
        "left_indent",
        "right_indent",
        "spacing_before",
        "spacing_after",
        "line_spacing",
        "wrap_text",
        "vertical_alignment",
        "fill_color",
        "number_format",
    }
)
_DOCX_ALIGNMENT_MAP = {
    "left": None if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.LEFT,
    "center": None if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.CENTER,
    "right": None if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": None if WD_ALIGN_PARAGRAPH is None else WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_DOCX_HIGHLIGHT_MAP = {
    "yellow": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.YELLOW,
    "green": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.BRIGHT_GREEN,
    "turquoise": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.TURQUOISE,
    "pink": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.PINK,
    "blue": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.BLUE,
    "red": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.RED,
    "dark_blue": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.DARK_BLUE,
    "teal": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.TEAL,
    "green_dark": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.GREEN,
    "violet": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.VIOLET,
    "dark_red": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.DARK_RED,
    "dark_yellow": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.DARK_YELLOW,
    "gray_50": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.GRAY_50,
    "gray_25": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.GRAY_25,
    "black": None if WD_COLOR_INDEX is None else WD_COLOR_INDEX.BLACK,
}
_DOCX_HIGHLIGHT_NAMES = {
    value: key for key, value in _DOCX_HIGHLIGHT_MAP.items() if value is not None
}


def _canonical_docx_locator(locator: str) -> tuple[str, tuple[str, ...]]:
    canonical = to_v2_locator(locator, file_type="docx")
    parsed = parse_locator(canonical)
    return canonical, parsed.components


def _normalize_clear_fields(
    clear_fields: list[str] | tuple[str, ...],
    allowed: frozenset[str],
) -> tuple[str, ...]:
    normalized: list[str] = []
    seen: set[str] = set()
    for field_name in clear_fields:
        if field_name not in allowed:
            raise InvalidArgumentsError(f"Unknown style field in clear_fields: {field_name}")
        if field_name not in seen:
            normalized.append(field_name)
            seen.add(field_name)
    return tuple(normalized)


def _apply_docx_inline_style(run: Run, style: InlineStyle, clear_fields: tuple[str, ...]) -> None:
    clear_set = set(clear_fields)
    if "bold" in clear_set:
        run.bold = None
    elif style.bold is not None:
        run.bold = style.bold

    if "italic" in clear_set:
        run.italic = None
    elif style.italic is not None:
        run.italic = style.italic

    if "underline" in clear_set:
        run.underline = None
    elif style.underline is not None:
        run.underline = style.underline

    if "strike" in clear_set:
        run.font.strike = None
    elif style.strike is not None:
        run.font.strike = style.strike

    if "font_name" in clear_set:
        run.font.name = None
    elif style.font_name is not None:
        run.font.name = style.font_name

    if "font_size" in clear_set:
        run.font.size = None
    elif style.font_size is not None:
        if Pt is None:
            raise RuntimeError("python-docx is required for DOCX operations.")
        run.font.size = Pt(style.font_size)

    if "font_color" in clear_set:
        run.font.color.rgb = None
    elif style.font_color is not None:
        run.font.color.rgb = RGBColor.from_string(_normalize_hex_color(style.font_color))

    if "highlight" in clear_set:
        run.font.highlight_color = None
    elif style.highlight is not None:
        run.font.highlight_color = _docx_highlight_value(style.highlight)


def _apply_docx_block_style(
    paragraph,
    style: BlockStyle,
    clear_fields: tuple[str, ...],
) -> list[str]:
    paragraph_format = paragraph.paragraph_format
    clear_set = set(clear_fields)
    skipped_fields: list[str] = []

    if "alignment" in clear_set:
        paragraph.alignment = None
    elif style.alignment is not None:
        paragraph.alignment = _docx_alignment_value(style.alignment)

    if "left_indent" in clear_set:
        paragraph_format.left_indent = None
    elif style.left_indent is not None:
        paragraph_format.left_indent = Pt(style.left_indent)

    if "right_indent" in clear_set:
        paragraph_format.right_indent = None
    elif style.right_indent is not None:
        paragraph_format.right_indent = Pt(style.right_indent)

    if "spacing_before" in clear_set:
        paragraph_format.space_before = None
    elif style.spacing_before is not None:
        paragraph_format.space_before = Pt(style.spacing_before)

    if "spacing_after" in clear_set:
        paragraph_format.space_after = None
    elif style.spacing_after is not None:
        paragraph_format.space_after = Pt(style.spacing_after)

    if "line_spacing" in clear_set:
        paragraph_format.line_spacing = None
    elif style.line_spacing is not None:
        paragraph_format.line_spacing = style.line_spacing

    for field_name in (
        "indent_level",
        "wrap_text",
        "vertical_alignment",
        "fill_color",
        "number_format",
    ):
        if getattr(style, field_name) is not None or field_name in clear_set:
            skipped_fields.append(field_name)

    return skipped_fields


def _docx_alignment_value(raw: str):
    normalized = raw.strip().lower()
    if normalized not in _DOCX_ALIGNMENT_MAP:
        raise InvalidArgumentsError(f"Unsupported DOCX alignment: {raw}")
    return _DOCX_ALIGNMENT_MAP[normalized]


def _docx_highlight_value(raw: str):
    normalized = raw.strip().lower()
    if normalized not in _DOCX_HIGHLIGHT_MAP:
        raise InvalidArgumentsError(f"Unsupported DOCX highlight color: {raw}")
    return _DOCX_HIGHLIGHT_MAP[normalized]


def _docx_highlight_name(value) -> str | None:
    return _DOCX_HIGHLIGHT_NAMES.get(value)


def _docx_structural_style_name(role: str, level: int | None) -> str:
    normalized = role.strip().lower()
    if normalized == "heading":
        if level is None or level < 1 or level > 9:
            raise InvalidArgumentsError("Heading structural role requires level between 1 and 9.")
        return f"Heading {level}"
    mapping = {
        "title": "Title",
        "body": "Normal",
        "table_header": "Table Heading",
        "caption": "Caption",
    }
    if normalized not in mapping:
        raise InvalidArgumentsError(f"Unsupported structural role: {role}")
    return mapping[normalized]


def _normalize_hex_color(value: str) -> str:
    normalized = value.strip().lstrip("#").upper()
    if len(normalized) != 6 or any(character not in "0123456789ABCDEF" for character in normalized):
        raise InvalidArgumentsError(f"Invalid RGB hex color: {value}")
    return normalized


def _parse_int_component(raw: str, locator: str) -> int:
    try:
        return int(raw)
    except ValueError as exc:
        raise InvalidArgumentsError(f"Invalid DOCX locator: {locator}") from exc
