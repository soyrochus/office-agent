from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from offagent.adapters import docx_adapter
from offagent.domain.locators import parse_locator, to_v2_locator
from offagent.domain.models import Capability, ChildSummary, ObjectPayload
from offagent.errors import InvalidArgumentsError, TargetNotFoundError

try:
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph
except ModuleNotFoundError:  # pragma: no cover - exercised through dependency checks
    WD_BREAK = None
    qn = None
    DocxParagraph = None


@dataclass(frozen=True)
class _DocxTarget:
    canonical_locator: str
    object_type: str
    paragraph_index: int | None = None
    run_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    column_index: int | None = None
    section_index: int | None = None
    image_index: int | None = None
    page_break_index: int | None = None


class DocxObjectResolver:
    def get_object(self, document_path: Path, locator: str) -> ObjectPayload:
        canonical = to_v2_locator(locator, file_type="docx")
        document = docx_adapter._open_document(document_path)
        target = _parse_docx_target(canonical)

        if target.object_type == "document":
            return _build_document_payload(document_path, document)
        if target.object_type == "section":
            return _build_section_payload(document_path, document, target)
        if target.object_type == "paragraph":
            return _build_paragraph_payload(document_path, document, target)
        if target.object_type == "run":
            return _build_run_payload(document_path, document, target)
        if target.object_type == "table":
            return _build_table_payload(document_path, document, target)
        if target.object_type == "table_row":
            return _build_table_row_payload(document_path, document, target)
        if target.object_type == "table_cell":
            return _build_table_cell_payload(document_path, document, target)
        if target.object_type == "image":
            return _build_image_payload(document_path, document, target)
        if target.object_type == "page_break":
            return _build_page_break_payload(document_path, document, target)
        raise InvalidArgumentsError(f"Unsupported DOCX object locator: {locator}")

    def list_children(
        self,
        document_path: Path,
        locator: str,
        *,
        child_type: str | None = None,
        limit: int | None = None,
    ) -> list[ChildSummary]:
        canonical = to_v2_locator(locator, file_type="docx")
        document = docx_adapter._open_document(document_path)
        target = _parse_docx_target(canonical)

        if target.object_type == "document":
            children = _document_children(document, child_type=child_type)
        elif target.object_type == "section":
            children = _section_children(document, target, child_type=child_type)
        elif target.object_type == "paragraph":
            children = _paragraph_children(document, target, child_type=child_type)
        elif target.object_type == "table":
            children = _table_children(document, target, child_type=child_type)
        elif target.object_type == "table_row":
            children = _table_row_children(document, target, child_type=child_type)
        else:
            children = ()

        if limit is not None:
            return list(children[:limit])
        return list(children)

    def resolve_capabilities(
        self, document_path: Path, locator: str
    ) -> frozenset[Capability]:
        del document_path
        canonical = to_v2_locator(locator, file_type="docx")
        target = _parse_docx_target(canonical)
        return _capabilities_for(target.object_type)


def set_paragraph_style(
    document_path: Path,
    locator: str,
    style_name: str,
    *,
    output_path: Path,
) -> tuple[str, str, dict[str, Any]]:
    canonical = to_v2_locator(locator, file_type="docx")
    document = docx_adapter._open_document(document_path)
    target = _parse_docx_target(canonical)
    if target.object_type != "paragraph":
        raise InvalidArgumentsError(
            "docx_set_paragraph_style requires a paragraph locator."
        )

    _require_style_name(document, style_name)
    resolved = _resolve_paragraph_target(document, target)
    resolved["paragraph"].style = style_name
    document.save(output_path)
    return (
        canonical,
        f"Applied paragraph style {style_name!r} to {canonical}.",
        {"style_name": style_name},
    )


def insert_page_break(
    document_path: Path,
    locator: str,
    *,
    output_path: Path,
) -> tuple[str, str, dict[str, Any]]:
    if DocxParagraph is None or WD_BREAK is None:
        raise RuntimeError("python-docx is required for DOCX operations.")

    canonical = to_v2_locator(locator, file_type="docx")
    document = docx_adapter._open_document(document_path)
    target = _parse_docx_target(canonical)
    if target.object_type != "paragraph":
        raise InvalidArgumentsError(
            "docx_insert_page_break requires a paragraph locator."
        )

    resolved = _resolve_paragraph_target(document, target)
    new_element = document.element.body.add_p()
    resolved["paragraph"]._element.addnext(new_element)
    paragraph = DocxParagraph(new_element, document)
    paragraph.add_run().add_break(WD_BREAK.PAGE)

    page_break = _page_break_for_paragraph(document, resolved["paragraph_index"] + 1)
    document.save(output_path)
    return (
        f"docx:page_break:{page_break['page_break_index']}",
        f"Inserted page break after {canonical}.",
        {"paragraph_locator": f"docx:para:{page_break['paragraph_index']}"},
    )


def add_table(
    document_path: Path,
    row_count: int,
    column_count: int,
    *,
    position: object | None = None,
    column_widths: list[int] | None = None,
    style_name: str | None = None,
    output_path: Path,
) -> tuple[str, str, dict[str, Any]]:
    if row_count < 1 or column_count < 1:
        raise InvalidArgumentsError(
            "docx_add_table requires positive row and column counts."
        )

    document = docx_adapter._open_document(document_path)
    table = document.add_table(rows=row_count, cols=column_count)

    if style_name is not None:
        _require_style_name(document, style_name)
        table.style = style_name

    if column_widths is not None:
        if len(column_widths) != column_count:
            raise InvalidArgumentsError(
                "column_widths must match the DOCX table column count."
            )
        for row in table.rows:
            for width, cell in zip(column_widths, row.cells, strict=True):
                cell.width = int(width)

    anchor = _resolve_insert_anchor(document, position)
    if anchor is not None:
        anchor.addnext(table._element)

    table_index = _table_index_for_element(document, table._element)
    locator = f"docx:table:{table_index}"
    document.save(output_path)
    return (
        locator,
        f"Inserted DOCX table {locator} with {row_count} rows and {column_count} columns.",
        {
            "row_count": row_count,
            "column_count": column_count,
            "style_name": style_name,
            "position": None if position is None else str(position),
        },
    )


def merge_table_cells(
    document_path: Path,
    start_locator: str,
    end_locator: str,
    *,
    output_path: Path,
) -> tuple[str, str, dict[str, Any]]:
    canonical_start = to_v2_locator(start_locator, file_type="docx")
    canonical_end = to_v2_locator(end_locator, file_type="docx")
    start_target = _parse_docx_target(canonical_start)
    end_target = _parse_docx_target(canonical_end)
    if (
        start_target.object_type != "table_cell"
        or end_target.object_type != "table_cell"
    ):
        raise InvalidArgumentsError(
            "docx_merge_table_cells requires DOCX table-cell locators."
        )
    if start_target.table_index != end_target.table_index:
        raise InvalidArgumentsError(
            "docx_merge_table_cells requires both cells to be in the same table."
        )

    document = docx_adapter._open_document(document_path)
    start = _resolve_table_cell_target(document, start_target)
    end = _resolve_table_cell_target(document, end_target)
    min_row, max_row = sorted((start["row_index"], end["row_index"]))
    min_col, max_col = sorted((start["column_index"], end["column_index"]))
    if min_row > max_row or min_col > max_col:
        raise InvalidArgumentsError(
            "Table-cell locators must define a valid rectangular range."
        )

    start["table"].cell(min_row, min_col).merge(start["table"].cell(max_row, max_col))
    document.save(output_path)
    locator = f"docx:table:{start['table_index']}:row:{min_row}:cell:{min_col}"
    return (
        locator,
        f"Merged DOCX table cells from {canonical_start} to {canonical_end}.",
        {"start_locator": canonical_start, "end_locator": canonical_end},
    )


def _build_document_payload(document_path: Path, document) -> ObjectPayload:
    blocks = docx_adapter._iter_blocks(document)
    paragraph_count = sum(1 for block_type, _ in blocks if block_type == "paragraph")
    table_count = sum(1 for block_type, _ in blocks if block_type == "table")
    preview = next(
        (
            block.text[:120]
            for block_type, block in blocks
            if block_type == "paragraph" and block.text.strip()
        ),
        "",
    )
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator="docx:document",
        object_type="document",
        preview=preview,
        properties={
            "section_count": len(document.sections),
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "image_count": len(document.inline_shapes),
            "page_break_count": len(_page_breaks(document)),
        },
        capabilities=_capability_tuple("document"),
        child_summary=_document_children(document),
    )


def _build_section_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    assert target.section_index is not None
    try:
        section = document.sections[target.section_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Section {target.section_index} does not exist in the document."
        ) from exc

    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="section",
        preview=f"Section {target.section_index}",
        properties={
            "section_index": target.section_index,
            "start_type": str(section.start_type),
            "page_width": int(section.page_width),
            "page_height": int(section.page_height),
            "left_margin": int(section.left_margin),
            "right_margin": int(section.right_margin),
        },
        capabilities=_capability_tuple("section"),
        parent_locator="docx:document",
        child_summary=_section_children(document, target),
    )


def _build_paragraph_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    resolved = _resolve_paragraph_target(document, target)
    paragraph_model = docx_adapter._paragraph_model(
        resolved["paragraph"],
        resolved["block_index"],
        resolved["paragraph_index"],
    )
    runs = tuple(docx_adapter._run_model(run) for run in resolved["paragraph"].runs)
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="paragraph",
        preview=paragraph_model.preview,
        properties={
            "block_index": resolved["block_index"],
            "paragraph_index": paragraph_model.paragraph_index,
            "text": paragraph_model.text,
            "style_name": paragraph_model.style_name,
            "is_heading": paragraph_model.is_heading,
            "runs": [
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "strike": run.strike,
                    "font_name": run.font_name,
                    "font_size": run.font_size,
                    "color_rgb": run.color_rgb,
                }
                for run in runs
            ],
        },
        capabilities=_capability_tuple("paragraph"),
        parent_locator="docx:document",
        child_summary=_paragraph_children(document, target),
    )


def _build_run_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    resolved = _resolve_run_target(document, target)
    run_model = docx_adapter._run_model(resolved["run"])
    paragraph_locator = f"docx:para:{resolved['paragraph_index']}"
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="run",
        preview=run_model.text[:120],
        properties={
            "paragraph_index": resolved["paragraph_index"],
            "run_index": resolved["run_index"],
            "text": run_model.text,
            "bold": run_model.bold,
            "italic": run_model.italic,
            "underline": run_model.underline,
            "strike": run_model.strike,
            "font_name": run_model.font_name,
            "font_size": run_model.font_size,
            "color_rgb": run_model.color_rgb,
        },
        capabilities=_capability_tuple("run"),
        parent_locator=paragraph_locator,
    )


def _build_table_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    resolved = _resolve_table_target(document, target)
    table_model = docx_adapter._table_model(
        resolved["table"],
        resolved["block_index"],
        resolved["table_index"],
    )
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="table",
        preview=table_model.preview,
        properties={
            "block_index": resolved["block_index"],
            "table_index": table_model.table_index,
            "row_count": len(table_model.rows),
            "column_count": max((len(row) for row in table_model.rows), default=0),
            "rows": [list(row) for row in table_model.rows],
        },
        capabilities=_capability_tuple("table"),
        parent_locator="docx:document",
        child_summary=_table_children(document, target),
    )


def _build_table_row_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    resolved = _resolve_table_row_target(document, target)
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="table_row",
        preview=" | ".join(cell.text for cell in resolved["row"].cells)[:120],
        properties={
            "table_index": resolved["table_index"],
            "row_index": resolved["row_index"],
            "cell_count": len(resolved["row"].cells),
            "cells": [cell.text for cell in resolved["row"].cells],
        },
        capabilities=_capability_tuple("table_row"),
        parent_locator=f"docx:table:{resolved['table_index']}",
        child_summary=_table_row_children(document, target),
    )


def _build_table_cell_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    resolved = _resolve_table_cell_target(document, target)
    cell = resolved["cell"]
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="table_cell",
        preview=cell.text[:120],
        properties={
            "table_index": resolved["table_index"],
            "row_index": resolved["row_index"],
            "column_index": resolved["column_index"],
            "text": cell.text,
            "paragraph_count": len(cell.paragraphs),
        },
        capabilities=_capability_tuple("table_cell"),
        parent_locator=f"docx:table:{resolved['table_index']}:row:{resolved['row_index']}",
    )


def _build_image_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    assert target.image_index is not None
    try:
        shape = document.inline_shapes[target.image_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Image {target.image_index} does not exist in the document."
        ) from exc

    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="image",
        preview=f"Image {target.image_index}",
        properties={
            "image_index": target.image_index,
            "width": int(shape.width),
            "height": int(shape.height),
            "shape_type": str(shape.type),
        },
        capabilities=_capability_tuple("image"),
        parent_locator="docx:document",
    )


def _build_page_break_payload(
    document_path: Path, document, target: _DocxTarget
) -> ObjectPayload:
    assert target.page_break_index is not None
    page_break = _resolve_page_break(document, target.page_break_index)
    return ObjectPayload(
        document=docx_adapter._document_ref(document_path),
        locator=target.canonical_locator,
        object_type="page_break",
        preview="Page break",
        properties={
            "page_break_index": target.page_break_index,
            "paragraph_index": page_break["paragraph_index"],
            "run_index": page_break["run_index"],
        },
        capabilities=_capability_tuple("page_break"),
        parent_locator=f"docx:para:{page_break['paragraph_index']}",
    )


def _document_children(
    document, *, child_type: str | None = None
) -> tuple[ChildSummary, ...]:
    children: list[ChildSummary] = []
    normalized_child_type = _normalize_child_type(child_type)
    if normalized_child_type in {None, "section"}:
        for section_index, _ in enumerate(document.sections):
            children.append(
                ChildSummary(
                    locator=f"docx:section:{section_index}",
                    object_type="section",
                    preview=f"Section {section_index}",
                    capabilities=_capability_tuple("section"),
                )
            )

    paragraph_index = 0
    table_index = 0
    for block_index, (block_type, block) in enumerate(
        docx_adapter._iter_blocks(document)
    ):
        if block_type == "paragraph":
            if normalized_child_type not in {None, "paragraph"}:
                paragraph_index += 1
                continue
            paragraph_model = docx_adapter._paragraph_model(
                block, block_index, paragraph_index
            )
            children.append(
                ChildSummary(
                    locator=f"docx:para:{paragraph_index}",
                    object_type="paragraph",
                    preview=paragraph_model.preview,
                    capabilities=_capability_tuple("paragraph"),
                )
            )
            paragraph_index += 1
            continue

        if normalized_child_type not in {None, "table"}:
            table_index += 1
            continue
        table_model = docx_adapter._table_model(block, block_index, table_index)
        children.append(
            ChildSummary(
                locator=f"docx:table:{table_index}",
                object_type="table",
                preview=table_model.preview,
                capabilities=_capability_tuple("table"),
            )
        )
        table_index += 1

    if normalized_child_type in {None, "image"}:
        for image_index, _ in enumerate(document.inline_shapes):
            children.append(
                ChildSummary(
                    locator=f"docx:image:{image_index}",
                    object_type="image",
                    preview=f"Image {image_index}",
                    capabilities=_capability_tuple("image"),
                )
            )

    if normalized_child_type in {None, "page_break"}:
        for page_break_index, _ in enumerate(_page_breaks(document)):
            children.append(
                ChildSummary(
                    locator=f"docx:page_break:{page_break_index}",
                    object_type="page_break",
                    preview="Page break",
                    capabilities=_capability_tuple("page_break"),
                )
            )

    return tuple(children)


def _section_children(
    document, target: _DocxTarget, *, child_type: str | None = None
) -> tuple[ChildSummary, ...]:
    if len(document.sections) == 1 and target.section_index == 0:
        return _document_children(document, child_type=child_type)
    return ()


def _paragraph_children(
    document, target: _DocxTarget, *, child_type: str | None = None
) -> tuple[ChildSummary, ...]:
    resolved = _resolve_paragraph_target(document, target)
    normalized_child_type = _normalize_child_type(child_type)
    children: list[ChildSummary] = []

    if normalized_child_type in {None, "run"}:
        for run_index, run in enumerate(resolved["paragraph"].runs):
            children.append(
                ChildSummary(
                    locator=f"docx:para:{resolved['paragraph_index']}:run:{run_index}",
                    object_type="run",
                    preview=run.text[:120],
                    capabilities=_capability_tuple("run"),
                )
            )

    if normalized_child_type in {None, "page_break"}:
        for page_break in _page_breaks_in_paragraph(
            resolved["paragraph"], resolved["paragraph_index"]
        ):
            children.append(
                ChildSummary(
                    locator=f"docx:page_break:{page_break['page_break_index']}",
                    object_type="page_break",
                    preview="Page break",
                    capabilities=_capability_tuple("page_break"),
                )
            )

    return tuple(children)


def _table_children(
    document, target: _DocxTarget, *, child_type: str | None = None
) -> tuple[ChildSummary, ...]:
    resolved = _resolve_table_target(document, target)
    normalized_child_type = _normalize_child_type(child_type)
    if normalized_child_type not in {None, "table_row"}:
        return ()
    return tuple(
        ChildSummary(
            locator=f"docx:table:{resolved['table_index']}:row:{row_index}",
            object_type="table_row",
            preview=" | ".join(cell.text for cell in row.cells)[:120],
            capabilities=_capability_tuple("table_row"),
        )
        for row_index, row in enumerate(resolved["table"].rows)
    )


def _table_row_children(
    document, target: _DocxTarget, *, child_type: str | None = None
) -> tuple[ChildSummary, ...]:
    resolved = _resolve_table_row_target(document, target)
    normalized_child_type = _normalize_child_type(child_type)
    if normalized_child_type not in {None, "table_cell"}:
        return ()
    return tuple(
        ChildSummary(
            locator=f"docx:table:{resolved['table_index']}:row:{resolved['row_index']}:cell:{column_index}",
            object_type="table_cell",
            preview=cell.text[:120],
            capabilities=_capability_tuple("table_cell"),
        )
        for column_index, cell in enumerate(resolved["row"].cells)
    )


def _resolve_paragraph_target(document, target: _DocxTarget) -> dict[str, Any]:
    assert target.paragraph_index is not None
    current_paragraph_index = 0
    for block_index, (block_type, block) in enumerate(
        docx_adapter._iter_blocks(document)
    ):
        if block_type != "paragraph":
            continue
        if current_paragraph_index == target.paragraph_index:
            return {
                "block_index": block_index,
                "paragraph_index": current_paragraph_index,
                "paragraph": block,
            }
        current_paragraph_index += 1
    raise TargetNotFoundError(
        f"Paragraph {target.paragraph_index} does not exist in the document."
    )


def _resolve_run_target(document, target: _DocxTarget) -> dict[str, Any]:
    resolved = _resolve_paragraph_target(document, target)
    assert target.run_index is not None
    try:
        run = resolved["paragraph"].runs[target.run_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Run {target.run_index} does not exist in paragraph {resolved['paragraph_index']}."
        ) from exc
    return {**resolved, "run_index": target.run_index, "run": run}


def _resolve_table_target(document, target: _DocxTarget) -> dict[str, Any]:
    assert target.table_index is not None
    current_table_index = 0
    for block_index, (block_type, block) in enumerate(
        docx_adapter._iter_blocks(document)
    ):
        if block_type != "table":
            continue
        if current_table_index == target.table_index:
            return {
                "block_index": block_index,
                "table_index": current_table_index,
                "table": block,
            }
        current_table_index += 1
    raise TargetNotFoundError(
        f"Table {target.table_index} does not exist in the document."
    )


def _resolve_table_row_target(document, target: _DocxTarget) -> dict[str, Any]:
    resolved = _resolve_table_target(document, target)
    assert target.row_index is not None
    try:
        row = resolved["table"].rows[target.row_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Row {target.row_index} does not exist in table {resolved['table_index']}."
        ) from exc
    return {**resolved, "row_index": target.row_index, "row": row}


def _resolve_table_cell_target(document, target: _DocxTarget) -> dict[str, Any]:
    resolved = _resolve_table_row_target(document, target)
    assert target.column_index is not None
    try:
        cell = resolved["row"].cells[target.column_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Cell {target.column_index} does not exist in row {resolved['row_index']}."
        ) from exc
    return {**resolved, "column_index": target.column_index, "cell": cell}


def _resolve_page_break(document, page_break_index: int) -> dict[str, int]:
    try:
        return _page_breaks(document)[page_break_index]
    except IndexError as exc:
        raise TargetNotFoundError(
            f"Page break {page_break_index} does not exist in the document."
        ) from exc


def _page_breaks(document) -> list[dict[str, int]]:
    page_breaks: list[dict[str, int]] = []
    paragraph_index = 0
    for block_type, block in docx_adapter._iter_blocks(document):
        if block_type != "paragraph":
            continue
        page_breaks.extend(
            _page_breaks_in_paragraph(
                block, paragraph_index, base_index=len(page_breaks)
            )
        )
        paragraph_index += 1
    return page_breaks


def _page_breaks_in_paragraph(
    paragraph, paragraph_index: int, *, base_index: int = 0
) -> list[dict[str, int]]:
    if qn is None:
        return []

    page_breaks: list[dict[str, int]] = []
    for run_index, run in enumerate(paragraph.runs):
        for br in run._element.findall(".//w:br", run._element.nsmap):
            if br.get(qn("w:type")) != "page":
                continue
            page_breaks.append(
                {
                    "page_break_index": base_index + len(page_breaks),
                    "paragraph_index": paragraph_index,
                    "run_index": run_index,
                }
            )
    return page_breaks


def _page_break_for_paragraph(document, paragraph_index: int) -> dict[str, int]:
    for page_break in _page_breaks(document):
        if page_break["paragraph_index"] == paragraph_index:
            return page_break
    raise RuntimeError(
        f"Failed to resolve inserted page break for paragraph {paragraph_index}."
    )


def _parse_docx_target(locator: str) -> _DocxTarget:
    parsed = parse_locator(locator)
    components = parsed.components
    if components == ("docx", "document"):
        return _DocxTarget(locator, "document")
    if len(components) == 3 and components[:2] == ("docx", "section"):
        return _DocxTarget(
            locator, "section", section_index=_require_index(components[2], locator)
        )
    if len(components) == 3 and components[:2] == ("docx", "para"):
        return _DocxTarget(
            locator, "paragraph", paragraph_index=_require_index(components[2], locator)
        )
    if (
        len(components) == 5
        and components[:2] == ("docx", "para")
        and components[3] == "run"
    ):
        return _DocxTarget(
            locator,
            "run",
            paragraph_index=_require_index(components[2], locator),
            run_index=_require_index(components[4], locator),
        )
    if len(components) == 3 and components[:2] == ("docx", "table"):
        return _DocxTarget(
            locator, "table", table_index=_require_index(components[2], locator)
        )
    if (
        len(components) == 5
        and components[:2] == ("docx", "table")
        and components[3] == "row"
    ):
        return _DocxTarget(
            locator,
            "table_row",
            table_index=_require_index(components[2], locator),
            row_index=_require_index(components[4], locator),
        )
    if (
        len(components) == 7
        and components[:2] == ("docx", "table")
        and components[3] == "row"
        and components[5] == "cell"
    ):
        return _DocxTarget(
            locator,
            "table_cell",
            table_index=_require_index(components[2], locator),
            row_index=_require_index(components[4], locator),
            column_index=_require_index(components[6], locator),
        )
    if len(components) == 3 and components[:2] == ("docx", "image"):
        return _DocxTarget(
            locator, "image", image_index=_require_index(components[2], locator)
        )
    if len(components) == 3 and components[:2] == ("docx", "page_break"):
        return _DocxTarget(
            locator,
            "page_break",
            page_break_index=_require_index(components[2], locator),
        )
    raise InvalidArgumentsError(f"Unsupported DOCX locator: {locator}")


def _capabilities_for(object_type: str) -> frozenset[Capability]:
    if object_type == "document":
        return frozenset({Capability.READ, Capability.ADD_CHILD})
    if object_type == "section":
        return frozenset({Capability.READ, Capability.ADD_CHILD})
    if object_type == "paragraph":
        return frozenset(
            {
                Capability.READ,
                Capability.UPDATE,
                Capability.DELETE,
                Capability.MOVE,
                Capability.COPY,
                Capability.STYLE,
            }
        )
    if object_type == "run":
        return frozenset(
            {Capability.READ, Capability.UPDATE, Capability.DELETE, Capability.STYLE}
        )
    if object_type == "table":
        return frozenset(
            {Capability.READ, Capability.DELETE, Capability.MOVE, Capability.COPY}
        )
    if object_type == "table_row":
        return frozenset(
            {
                Capability.READ,
                Capability.UPDATE,
                Capability.DELETE,
                Capability.ADD_CHILD,
                Capability.MOVE,
                Capability.COPY,
            }
        )
    if object_type == "table_cell":
        return frozenset({Capability.READ, Capability.UPDATE, Capability.STYLE})
    if object_type in {"image", "page_break"}:
        return frozenset(
            {Capability.READ, Capability.DELETE, Capability.MOVE, Capability.COPY}
        )
    return frozenset({Capability.READ})


def _capability_tuple(object_type: str) -> tuple[Capability, ...]:
    return tuple(
        sorted(_capabilities_for(object_type), key=lambda capability: capability.value)
    )


def _normalize_child_type(child_type: str | None) -> str | None:
    if child_type in {None, ""}:
        return None
    return child_type


def _require_style_name(document, style_name: str) -> None:
    if any(getattr(style, "name", None) == style_name for style in document.styles):
        return
    raise InvalidArgumentsError(f"Unknown DOCX style: {style_name}")


def _resolve_insert_anchor(document, position: object | None):
    if position is None:
        return None

    after_locator: str | None = None
    if isinstance(position, str):
        after_locator = position
    elif isinstance(position, dict):
        for key in ("after", "after_locator"):
            value = position.get(key)
            if value is not None:
                after_locator = str(value)
                break
    if after_locator is None:
        raise InvalidArgumentsError("DOCX insert position must be an after locator.")

    canonical = to_v2_locator(after_locator, file_type="docx")
    target = _parse_docx_target(canonical)
    if target.object_type == "paragraph":
        return _resolve_paragraph_target(document, target)["paragraph"]._element
    if target.object_type == "table":
        return _resolve_table_target(document, target)["table"]._element
    raise InvalidArgumentsError(
        "DOCX insert position must reference a paragraph or table."
    )


def _table_index_for_element(document, table_element) -> int:
    table_index = 0
    for block_type, block in docx_adapter._iter_blocks(document):
        if block_type != "table":
            continue
        if block._element == table_element:
            return table_index
        table_index += 1
    raise RuntimeError("Failed to resolve inserted DOCX table index.")


def _require_index(raw: str, locator: str) -> int:
    try:
        return int(raw)
    except ValueError as exc:
        raise InvalidArgumentsError(f"Invalid DOCX locator: {locator}") from exc
