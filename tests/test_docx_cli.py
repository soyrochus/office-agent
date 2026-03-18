from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document


def _run_cli(config_path, *args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [
            sys.executable,
            "-m",
            "offagent",
            *args,
            "--config",
            str(config_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )


def _parse_output_path(stdout: str) -> Path:
    return Path(stdout.strip().split("\t")[-1])


def test_docx_cli_round_trip(sample_docx, config_path) -> None:
    index_result = _run_cli(config_path, "index", str(sample_docx))
    search_result = _run_cli(config_path, "search", "Supplier shall", "--type", "docx")
    locate_result = _run_cli(config_path, "locate", "--doc", str(sample_docx), "--paragraph", "3")
    read_result = _run_cli(config_path, "read", "--doc", str(sample_docx), "--item", "para:3")
    replace_result = _run_cli(
        config_path,
        "replace",
        "--doc",
        str(sample_docx),
        "--item",
        "para:1",
        "--text",
        "CLI replaced text.",
    )
    replace_output_path = _parse_output_path(replace_result.stdout)
    append_result = _run_cli(
        config_path,
        "append",
        "--doc",
        str(replace_output_path),
        "--item",
        "para:2",
        "--text",
        "CLI append text.",
    )
    append_output_path = _parse_output_path(append_result.stdout)

    assert index_result.returncode == 0, index_result.stderr
    assert "indexed 1" in index_result.stdout
    assert search_result.returncode == 0, search_result.stderr
    assert "para:3" in search_result.stdout
    assert locate_result.returncode == 0, locate_result.stderr
    assert "para:3" in locate_result.stdout
    assert read_result.returncode == 0, read_result.stderr
    assert "Supplier shall deliver by Friday." in read_result.stdout
    assert replace_result.returncode == 0, replace_result.stderr
    assert append_result.returncode == 0, append_result.stderr

    original_document = Document(str(sample_docx))
    final_document = Document(str(append_output_path))
    assert original_document.paragraphs[1].text == "Alpha paragraph for search."
    assert final_document.paragraphs[1].text == "CLI replaced text."
    assert final_document.paragraphs[1].runs[0].bold is True
    assert final_document.paragraphs[2].text == "CLI append text."


def test_docx_cli_supports_inplace_output_mode_when_enabled(sample_docx, tmp_path) -> None:
    config_path = tmp_path / "office-agent.toml"
    config_path.write_text(
        f"""
[offagent]
index_path = "{(tmp_path / 'state' / 'index.sqlite3').as_posix()}"
document_roots = ["{tmp_path.as_posix()}"]
allow_inplace_overwrite = true
""".strip()
    )

    _run_cli(config_path, "index", str(sample_docx))
    replace_result = _run_cli(
        config_path,
        "replace",
        "--doc",
        str(sample_docx),
        "--item",
        "para:1",
        "--text",
        "CLI inplace text.",
        "--output-mode",
        "inplace",
    )

    document = Document(str(sample_docx))
    assert replace_result.returncode == 0, replace_result.stderr
    assert str(sample_docx) in replace_result.stdout
    assert document.paragraphs[1].text == "CLI inplace text."
