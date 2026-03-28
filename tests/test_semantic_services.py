from __future__ import annotations

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from offagent.app.services import AppServices
from offagent.config import AppConfig
from offagent.errors import InvalidArgumentsError


def _services(tmp_path):
    edited_dir = tmp_path / "edited"
    return AppServices(
        AppConfig(
            index_path=tmp_path / "state" / "index.sqlite3",
            document_roots=(tmp_path,),
            output_directory=edited_dir,
            output_roots=(edited_dir, tmp_path),
        )
    )


def test_semantic_service_reads_across_formats(sample_docx, sample_pptx, sample_xlsx, tmp_path) -> None:
    services = _services(tmp_path)
    docx_document = services.index_document(sample_docx)
    pptx_document = services.index_document(sample_pptx)
    xlsx_document = services.index_document(sample_xlsx)

    docx_structure = services.get_structure(docx_document.document_id)
    pptx_structure = services.get_structure(pptx_document.document_id)
    xlsx_structure = services.get_structure(xlsx_document.document_id)
    docx_section = services.get_section(docx_document.document_id, "para:0")
    pptx_section = services.get_section(
        pptx_document.document_id,
        pptx_structure.sections[0].locator,
    )
    xlsx_section = services.get_section(
        xlsx_document.document_id,
        xlsx_structure.sections[1].locator,
        cell_range="A1:B2",
    )
    table_section = services.get_section(docx_document.document_id, "table:0:cell:0:0")
    docx_tables = services.docx_get_tables(docx_document.document_id)
    node = services.get_node(docx_document.document_id, "para:3")

    assert [section.section_type for section in docx_structure.sections] == [
        "paragraph",
        "paragraph",
        "paragraph",
        "paragraph",
        "table",
    ]
    assert [section.section_type for section in pptx_structure.sections] == ["slide", "slide"]
    assert [section.metadata["sheet_name"] for section in xlsx_structure.sections] == [
        "Budget2026",
        "Notes 2026",
    ]
    assert docx_section.text == "Project Heading"
    assert pptx_section.notes_text == "Speaker notes: confirm launch owner."
    assert [cell.coordinate for cell in xlsx_section.cells] == ["A1", "B1", "A2", "B2"]
    assert table_section.rows[0][0] == "Table text to ignore"
    assert docx_tables.tables[0].locator == "table:0:cell:0:0"
    assert node.text == "Supplier shall deliver by Friday."


def test_semantic_service_writes_return_reindexed_results(sample_docx, sample_xlsx, tmp_path) -> None:
    services = _services(tmp_path)
    docx_document = services.index_document(sample_docx)
    xlsx_document = services.index_document(sample_xlsx)

    write_result = services.write_node(
        docx_document.document_id,
        "para:1",
        "Semantic replacement paragraph.",
    )
    insert_result = services.insert_content(
        write_result.document_id,
        "Semantic appendix paragraph.",
        style_name="Heading 2",
        after_node_id="para:1",
    )
    append_rows_result = services.xlsx_insert_rows(
        xlsx_document.document_id,
        "Notes 2026",
        rows=[["Escalate owner", "Finance"]],
    )

    replaced_document = Document(str(write_result.output_path))
    inserted_document = Document(str(insert_result.output_path))
    appended_workbook = load_workbook(append_rows_result.output_path)

    assert write_result.previous_text == "Alpha paragraph for search."
    assert replaced_document.paragraphs[1].text == "Semantic replacement paragraph."
    assert insert_result.new_node_id == "para:2"
    assert inserted_document.paragraphs[2].text == "Semantic appendix paragraph."
    assert inserted_document.paragraphs[2].style.name == "Heading 2"
    assert append_rows_result.first_row_locator == "sheet:Notes 2026!A3"
    assert appended_workbook["Notes 2026"]["A3"].value == "Escalate owner"
    assert services.get_node(insert_result.document_id, insert_result.new_node_id).text == "Semantic appendix paragraph."


def test_semantic_service_rejects_unsupported_targets(sample_docx, sample_xlsx, tmp_path) -> None:
    services = _services(tmp_path)
    docx_document = services.index_document(sample_docx)
    xlsx_document = services.index_document(sample_xlsx)

    with pytest.raises(InvalidArgumentsError, match="requires a .xlsx document"):
        services.xlsx_insert_rows(docx_document.document_id, "Sheet1", rows=[["blocked"]])

    with pytest.raises(InvalidArgumentsError, match="requires a .docx document"):
        services.insert_content(xlsx_document.document_id, "blocked")

    with pytest.raises(InvalidArgumentsError, match="requires a .docx document"):
        services.docx_get_tables(xlsx_document.document_id)


def test_semantic_service_xlsx_insert_rows_supports_mapped_records(tmp_path) -> None:
    services = _services(tmp_path)
    workbook_path = tmp_path / "mapped.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Tasks"
    worksheet["A1"] = "Task"
    worksheet["B1"] = "Owner"
    workbook.save(workbook_path)
    document = services.index_document(workbook_path)

    result = services.xlsx_insert_rows(
        document.document_id,
        "Tasks",
        records=[
            {"Task": "Review budget", "Owner": "Finance"},
            {"Task": "Confirm launch", "Owner": "Operations"},
        ],
    )

    written_workbook = load_workbook(result.output_path)
    assert result.first_row_locator == "sheet:Tasks!A2"
    assert result.rows_inserted == 2
    assert written_workbook["Tasks"]["A2"].value == "Review budget"
    assert written_workbook["Tasks"]["B3"].value == "Operations"
