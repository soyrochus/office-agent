from __future__ import annotations

import pytest
from docx import Document

from offagent.app.services import AppServices
from offagent.config import AppConfig


def test_index_search_locate_and_read_docx(sample_docx, tmp_path) -> None:
    services = AppServices(
        AppConfig(index_path=tmp_path / "state" / "index.sqlite3", document_roots=(tmp_path,))
    )

    summary = services.index_path(sample_docx)
    hits = services.search_corpus("Supplier shall", file_type="docx")
    item = services.locate_paragraph(sample_docx, 3)
    text = services.read_item(sample_docx, "para:3")

    assert summary.files_indexed == 1
    assert hits[0].item_id == "para:3"
    assert hits[0].preview == "Supplier shall deliver by Friday."
    assert item.item_id == "para:3"
    assert text == "Supplier shall deliver by Friday."


def test_replace_append_and_reindex_docx(sample_docx, tmp_path) -> None:
    services = AppServices(
        AppConfig(index_path=tmp_path / "state" / "index.sqlite3", document_roots=(tmp_path,))
    )
    services.index_document(sample_docx)

    replace_result = services.replace_item_text(sample_docx, "para:1", "Replaced alpha text.")
    append_result = services.append_item_text(replace_result.output_path, "para:2", "Now populated.")

    original_document = Document(str(sample_docx))
    replaced_document = Document(str(replace_result.output_path))
    appended_document = Document(str(append_result.output_path))

    assert replace_result.text == "Replaced alpha text."
    assert append_result.text == "Now populated."
    assert replace_result.output_path != sample_docx
    assert ".edited." in replace_result.output_path.name
    assert original_document.paragraphs[1].text == "Alpha paragraph for search."
    assert replaced_document.paragraphs[1].text == "Replaced alpha text."
    assert appended_document.paragraphs[1].runs[0].bold is True
    assert appended_document.paragraphs[2].text == "Now populated."

    hits = services.search_corpus("Now populated")
    assert hits[0].item_id == "para:2"
    assert hits[0].document_path == append_result.output_path


def test_replace_inplace_is_allowed_by_default_docx(sample_docx, tmp_path) -> None:
    services = AppServices(
        AppConfig(index_path=tmp_path / "state" / "index.sqlite3", document_roots=(tmp_path,))
    )
    services.index_document(sample_docx)

    result = services.replace_item_text(
        sample_docx,
        "para:1",
        "Allowed overwrite.",
        output_mode="inplace",
    )

    document = Document(str(sample_docx))
    assert result.output_path == sample_docx
    assert document.paragraphs[1].text == "Allowed overwrite."


def test_replace_inplace_can_be_disabled_explicitly_docx(sample_docx, tmp_path) -> None:
    services = AppServices(
        AppConfig(
            index_path=tmp_path / "state" / "index.sqlite3",
            document_roots=(tmp_path,),
            allow_inplace_overwrite=False,
        )
    )
    services.index_document(sample_docx)

    with pytest.raises(RuntimeError, match="In-place overwrite is not enabled"):
        services.replace_item_text(sample_docx, "para:1", "Blocked overwrite.", output_mode="inplace")
