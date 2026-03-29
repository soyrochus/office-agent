from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from offagent.domain.models import Capability
from offagent.errors import InvalidArgumentsError
from offagent.objects.docx_objects import (
    DocxObjectResolver,
    add_table,
    insert_page_break,
    merge_table_cells,
    set_paragraph_style,
)

PNG_1X1 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/aRsAAAAASUVORK5CYII="
)


@pytest.fixture
def docx_with_page_break_and_image(tmp_path) -> Path:
    image_path = tmp_path / "dot.png"
    image_path.write_bytes(base64.b64decode(PNG_1X1))

    document = Document()
    paragraph = document.add_paragraph("Before break")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("After break")
    document.add_picture(str(image_path))

    path = tmp_path / "rich.docx"
    document.save(path)
    return path


@pytest.fixture
def docx_with_mergeable_table(tmp_path) -> Path:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    path = tmp_path / "mergeable.docx"
    document.save(path)
    return path


def test_get_document_and_paragraph_objects(sample_docx) -> None:
    resolver = DocxObjectResolver()

    document = resolver.get_object(sample_docx, "docx:document")
    paragraph = resolver.get_object(sample_docx, "para:1")

    assert document.object_type == "document"
    assert document.locator == "docx:document"
    assert document.properties["paragraph_count"] == 4
    assert Capability.ADD_CHILD in document.capabilities
    assert any(child.object_type == "paragraph" for child in document.child_summary)

    assert paragraph.object_type == "paragraph"
    assert paragraph.locator == "docx:para:1"
    assert paragraph.properties["text"] == "Alpha paragraph for search."
    assert paragraph.parent_locator == "docx:document"
    assert paragraph.child_summary[0].locator == "docx:para:1:run:0"


def test_get_table_row_and_cell_objects(sample_docx) -> None:
    resolver = DocxObjectResolver()

    table = resolver.get_object(sample_docx, "docx:table:0")
    row = resolver.get_object(sample_docx, "docx:table:0:row:0")
    cell = resolver.get_object(sample_docx, "table:0:cell:0:0")

    assert table.object_type == "table"
    assert table.properties["row_count"] == 1
    assert row.object_type == "table_row"
    assert row.child_summary[0].locator == "docx:table:0:row:0:cell:0"
    assert cell.object_type == "table_cell"
    assert cell.properties["text"] == "Table text to ignore"
    assert Capability.UPDATE in cell.capabilities


def test_list_children_filters_document_and_table_children(sample_docx) -> None:
    resolver = DocxObjectResolver()

    paragraphs = resolver.list_children(sample_docx, "docx:document", child_type="paragraph")
    rows = resolver.list_children(sample_docx, "docx:table:0", child_type="table_row")

    assert [child.locator for child in paragraphs] == [
        "docx:para:0",
        "docx:para:1",
        "docx:para:2",
        "docx:para:3",
    ]
    assert len(rows) == 1
    assert rows[0].locator == "docx:table:0:row:0"


def test_resolve_capabilities_for_docx_objects(sample_docx) -> None:
    resolver = DocxObjectResolver()

    paragraph_capabilities = resolver.resolve_capabilities(sample_docx, "docx:para:1")
    run_capabilities = resolver.resolve_capabilities(sample_docx, "docx:para:1:run:0")

    assert paragraph_capabilities == frozenset(
        {
            Capability.READ,
            Capability.UPDATE,
            Capability.DELETE,
            Capability.MOVE,
            Capability.COPY,
            Capability.STYLE,
        }
    )
    assert run_capabilities == frozenset(
        {
            Capability.READ,
            Capability.UPDATE,
            Capability.DELETE,
            Capability.STYLE,
        }
    )


def test_get_image_and_page_break_objects(docx_with_page_break_and_image) -> None:
    resolver = DocxObjectResolver()

    page_break = resolver.get_object(docx_with_page_break_and_image, "docx:page_break:0")
    image = resolver.get_object(docx_with_page_break_and_image, "docx:image:0")

    assert page_break.object_type == "page_break"
    assert page_break.parent_locator == "docx:para:0"
    assert image.object_type == "image"
    assert image.properties["image_index"] == 0


def test_docx_escape_hatches_apply_style_insert_break_add_table_and_merge(
    sample_docx,
    docx_with_mergeable_table,
    tmp_path,
) -> None:
    styled_path = tmp_path / "styled.docx"
    break_path = tmp_path / "break.docx"
    table_path = tmp_path / "table.docx"
    merged_path = tmp_path / "merged.docx"
    resolver = DocxObjectResolver()

    styled_locator, _, _ = set_paragraph_style(
        sample_docx,
        "docx:para:1",
        "Heading 2",
        output_path=styled_path,
    )
    page_break_locator, _, _ = insert_page_break(
        sample_docx,
        "docx:para:1",
        output_path=break_path,
    )
    table_locator, _, _ = add_table(
        sample_docx,
        2,
        2,
        style_name="Table Grid",
        output_path=table_path,
    )
    merged_locator, _, _ = merge_table_cells(
        docx_with_mergeable_table,
        "docx:table:0:row:0:cell:0",
        "docx:table:0:row:0:cell:1",
        output_path=merged_path,
    )

    assert styled_locator == "docx:para:1"
    assert Document(str(styled_path)).paragraphs[1].style.name == "Heading 2"

    page_break = resolver.get_object(break_path, page_break_locator)
    assert page_break.object_type == "page_break"
    assert page_break.parent_locator == "docx:para:2"

    added_table = resolver.get_object(table_path, table_locator)
    assert added_table.object_type == "table"
    assert added_table.properties["row_count"] == 2
    assert added_table.properties["column_count"] == 2
    assert Document(str(table_path)).tables[-1].style.name == "Table Grid"

    merged_table = Document(str(merged_path)).tables[0]
    assert merged_locator == "docx:table:0:row:0:cell:0"
    assert merged_table.cell(0, 0)._tc.tcPr.gridSpan.val == 2


def test_docx_set_paragraph_style_rejects_unknown_style(sample_docx, tmp_path) -> None:
    with pytest.raises(InvalidArgumentsError, match="Unknown DOCX style"):
        set_paragraph_style(
            sample_docx,
            "docx:para:1",
            "Missing Style",
            output_path=tmp_path / "missing-style.docx",
        )
