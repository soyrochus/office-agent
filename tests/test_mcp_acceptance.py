from __future__ import annotations

import asyncio
import sys
from pathlib import Path

from docx import Document
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client
from openpyxl import load_workbook


def _server_params(config_path: Path) -> StdioServerParameters:
    return StdioServerParameters(
        command=sys.executable,
        args=["-m", "offagent", "mcp", "--config", str(config_path)],
    )


def _run_mcp(config_path: Path, callback):
    async def runner():
        async with stdio_client(_server_params(config_path)) as (read_stream, write_stream):
            async with ClientSession(read_stream, write_stream) as session:
                await session.initialize()
                return await callback(session)

    return asyncio.run(runner())


async def _call_tool(session: ClientSession, name: str, arguments: dict | None = None):
    result = await session.call_tool(name, arguments=arguments)
    assert not result.isError, _tool_error_text(result)
    return result.structuredContent


def _tool_error_text(result) -> str:
    parts: list[str] = []
    for content in result.content:
        text = getattr(content, "text", None)
        if text is not None:
            parts.append(text)
    return "\n".join(parts)


def test_mcp_acceptance_surface(golden_config_path, golden_docx, golden_xlsx, golden_workspace) -> None:
    docs_dir = golden_workspace / "docs"

    def scenario(session: ClientSession):
        async def run():
            tools = await session.list_tools()
            tool_map = {tool.name: tool for tool in tools.tools}
            assert set(tool_map) == {
                "index_documents",
                "refresh_document",
                "list_documents",
                "search_objects",
                "search_documents",
                "get_structure",
                "get_section",
                "get_object",
                "list_children",
                "get_node",
                "write_node",
                "create_object",
                "update_object",
                "move_object",
                "copy_object",
                "batch_edit",
                "delete_object",
                "insert_content",
                "xlsx_insert_rows",
                "xlsx_write_range",
                "xlsx_insert_columns",
                "xlsx_set_formula",
                "xlsx_merge_cells",
                "docx_get_tables",
                "docx_set_paragraph_style",
                "docx_insert_page_break",
                "docx_add_table",
                "docx_merge_table_cells",
                "pptx_add_slide",
                "pptx_duplicate_slide",
                "pptx_set_slide_layout",
                "pptx_add_text_shape",
            }
            assert all(tool.inputSchema is not None for tool in tool_map.values())
            assert all(tool.outputSchema is not None for tool in tool_map.values())

            index_result = await _call_tool(
                session,
                "index_documents",
                {"paths": [str(docs_dir)]},
            )
            assert index_result["files_indexed"] == 3

            documents_result = await _call_tool(session, "list_documents")
            assert len(documents_result["documents"]) == 3
            docx_document = next(
                document for document in documents_result["documents"] if document["path"] == str(golden_docx)
            )
            xlsx_document = next(
                document for document in documents_result["documents"] if document["path"] == str(golden_xlsx)
            )
            assert docx_document["item_count"] == 11

            search_result = await _call_tool(
                session,
                "search_documents",
                {"query": "Supplier shall", "file_type": "docx"},
            )
            hit = search_result["hits"][0]
            assert hit["item_id"] == "para:3"

            read_result = await _call_tool(
                session,
                "get_node",
                {"document_id": docx_document["document_id"], "node_id": "para:3"},
            )
            assert read_result["text"] == "Supplier shall deliver by Friday."

            structure_result = await _call_tool(
                session,
                "get_structure",
                {"document_id": docx_document["document_id"]},
            )
            paragraph_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docx_document["document_id"],
                    "section_id": structure_result["sections"][4]["locator"],
                },
            )
            assert paragraph_section["text"] == "Timeline checkpoint: confirm rollout owners by Wednesday."

            replace_result = await _call_tool(
                session,
                "write_node",
                {
                    "document_id": docx_document["document_id"],
                    "node_id": "para:1",
                    "content": "MCP acceptance text.",
                },
            )
            append_result = await _call_tool(
                session,
                "write_node",
                {
                    "document_id": replace_result["document_id"],
                    "node_id": "para:2",
                    "content": "MCP appended acceptance text.",
                },
            )
            refresh_result = await _call_tool(
                session,
                "refresh_document",
                {"document_id": docx_document["document_id"]},
            )
            assert refresh_result["files_indexed"] == 1

            write_result = await _call_tool(
                session,
                "write_node",
                {
                    "document_id": xlsx_document["document_id"],
                    "node_id": "sheet:Budget2026!B2",
                    "content": "125001",
                },
            )
            append_rows = await _call_tool(
                session,
                "xlsx_insert_rows",
                {
                    "document_id": write_result["document_id"],
                    "sheet_name": "Notes 2026",
                    "rows": [["MCP acceptance note.", "Finance"]],
                },
            )

            return {
                "docx_output": append_result["output_path"],
                "xlsx_output": append_rows["output_path"],
            }

        return run()

    outputs = _run_mcp(golden_config_path, scenario)

    final_document = Document(outputs["docx_output"])
    assert final_document.paragraphs[1].text == "MCP acceptance text."
    assert final_document.paragraphs[2].text == "MCP appended acceptance text."

    workbook = load_workbook(outputs["xlsx_output"])
    assert workbook["Budget2026"]["B2"].value == 125001
    assert workbook["Notes 2026"]["A4"].value == "MCP acceptance note."


def test_mcp_acceptance_errors_and_policy_refusal(
    golden_config_path,
    golden_docx,
    golden_xlsx,
    golden_workspace,
    tmp_path,
) -> None:
    docs_dir = golden_workspace / "docs"

    def scenario(session: ClientSession):
        async def run():
            await _call_tool(session, "index_documents", {"paths": [str(docs_dir)]})
            documents_result = await _call_tool(session, "list_documents")
            docs_by_path = {document["path"]: document for document in documents_result["documents"]}

            unsupported = await session.call_tool(
                "insert_content",
                {"document_id": docs_by_path[str(golden_xlsx)]["document_id"], "content": "blocked"},
            )
            assert unsupported.isError
            assert "requires a .docx document" in _tool_error_text(unsupported)

            workbook = load_workbook(golden_xlsx)
            del workbook["Budget2026"]
            workbook.save(golden_xlsx)
            stale_result = await session.call_tool(
                "write_node",
                {
                    "document_id": docs_by_path[str(golden_xlsx)]["document_id"],
                    "node_id": "sheet:Budget2026!B2",
                    "content": "125001",
                },
            )
            assert stale_result.isError
            assert "stale locator" in _tool_error_text(stale_result)

        return run()

    _run_mcp(golden_config_path, scenario)

    policy_config = tmp_path / "policy-office-agent.toml"
    policy_config.write_text(
        f"""
[offagent]
index_path = "{(tmp_path / 'state' / 'index.sqlite3').as_posix()}"
document_roots = ["{docs_dir.as_posix()}"]
allowed_roots = ["{docs_dir.as_posix()}"]
output_directory = "{(tmp_path / 'edited').as_posix()}"
output_roots = ["{(tmp_path / 'elsewhere').as_posix()}"]
""".strip()
    )

    def policy_scenario(session: ClientSession):
        async def run():
            await _call_tool(session, "index_documents", {"paths": [str(golden_docx)]})
            documents_result = await _call_tool(session, "list_documents")
            document = documents_result["documents"][0]
            refused = await session.call_tool(
                "write_node",
                {
                    "document_id": document["document_id"],
                    "node_id": "para:1",
                    "content": "blocked",
                },
            )
            assert refused.isError
            assert "output roots" in _tool_error_text(refused)

        return run()

    _run_mcp(policy_config, policy_scenario)


def test_mcp_acceptance_semantic_surface(
    golden_config_path,
    golden_docx,
    golden_pptx,
    golden_xlsx,
    golden_workspace,
) -> None:
    docs_dir = golden_workspace / "docs"

    def scenario(session: ClientSession):
        async def run():
            tools = await session.list_tools()
            tool_names = {tool.name for tool in tools.tools}
            assert {
                "index_documents",
                "list_documents",
                "get_structure",
                "get_section",
                "search_documents",
                "search_objects",
                "get_object",
                "list_children",
            } <= tool_names

            await _call_tool(session, "index_documents", {"paths": [str(docs_dir)]})
            documents_result = await _call_tool(session, "list_documents")
            docs_by_path = {document["path"]: document for document in documents_result["documents"]}

            docx_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_path[str(golden_docx)]["document_id"]},
            )
            pptx_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_path[str(golden_pptx)]["document_id"]},
            )
            slide_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docs_by_path[str(golden_pptx)]["document_id"],
                    "section_id": pptx_structure["sections"][0]["locator"],
                },
            )
            workbook_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_path[str(golden_xlsx)]["document_id"]},
            )
            sheet_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docs_by_path[str(golden_xlsx)]["document_id"],
                    "section_id": workbook_structure["sections"][1]["locator"],
                    "cell_range": "A1:B2",
                },
            )
            tables = await _call_tool(
                session,
                "docx_get_tables",
                {"document_id": docs_by_path[str(golden_docx)]["document_id"]},
            )
            insert_content = await _call_tool(
                session,
                "insert_content",
                {
                    "document_id": docs_by_path[str(golden_docx)]["document_id"],
                    "content": "Golden semantic appendix.",
                    "style_name": "Heading 2",
                },
            )
            append_row = await _call_tool(
                session,
                "xlsx_insert_rows",
                {
                    "document_id": docs_by_path[str(golden_xlsx)]["document_id"],
                    "sheet_name": "Notes 2026",
                    "rows": [["Golden semantic note", "Operations"]],
                },
            )

            return {
                "docx_structure": docx_structure,
                "pptx_structure": pptx_structure,
                "slide_section": slide_section,
                "workbook_structure": workbook_structure,
                "sheet_section": sheet_section,
                "tables": tables,
                "insert_content": insert_content,
                "append_row": append_row,
            }

        return run()

    result = _run_mcp(golden_config_path, scenario)

    assert len(result["docx_structure"]["sections"]) == 11
    assert [slide["slide_number"] for slide in result["pptx_structure"]["sections"]] == [1, 2, 3]
    assert result["slide_section"]["slide_number"] == 1
    assert result["slide_section"]["notes_text"] == ""
    assert [sheet["sheet_name"] for sheet in result["workbook_structure"]["sections"]] == [
        "Budget2026",
        "Notes 2026",
    ]
    assert [cell["coordinate"] for cell in result["sheet_section"]["cells"]] == ["A1", "B1", "A2", "B2"]
    assert result["tables"]["tables"] == []

    final_document = Document(result["insert_content"]["output_path"])
    workbook = load_workbook(result["append_row"]["output_path"])
    assert final_document.paragraphs[-1].text == "Golden semantic appendix."
    assert workbook["Notes 2026"]["A4"].value == "Golden semantic note"
