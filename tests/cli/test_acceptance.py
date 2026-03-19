from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


def _run_cli(config_path: Path, *args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [
            sys.executable,
            "-m",
            "offagent",
            *args,
            "--config",
            str(config_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )


def _parse_output_path(stdout: str) -> Path:
    return Path(stdout.strip().split("\t")[-1])


def test_cli_acceptance_surface(golden_config_path, golden_docx, golden_xlsx, golden_workspace) -> None:
    docs_dir = golden_workspace / "docs"

    index_result = _run_cli(golden_config_path, "index", str(docs_dir))
    reindex_result = _run_cli(golden_config_path, "reindex", str(golden_xlsx))
    list_result = _run_cli(golden_config_path, "list")
    show_doc_result = _run_cli(golden_config_path, "show", "--doc", str(golden_docx))
    show_item_result = _run_cli(
        golden_config_path,
        "show",
        "--doc",
        str(golden_docx),
        "--item",
        "para:3",
    )
    search_result = _run_cli(golden_config_path, "search", "Supplier shall", "--type", "docx")
    locate_result = _run_cli(
        golden_config_path,
        "locate",
        "--doc",
        str(golden_docx),
        "--paragraph",
        "3",
    )
    read_result = _run_cli(
        golden_config_path,
        "read",
        "--doc",
        str(golden_docx),
        "--item",
        "para:3",
    )
    replace_result = _run_cli(
        golden_config_path,
        "replace",
        "--doc",
        str(golden_docx),
        "--item",
        "para:1",
        "--text",
        "CLI acceptance text.",
    )
    replace_output_path = _parse_output_path(replace_result.stdout)
    append_result = _run_cli(
        golden_config_path,
        "append",
        "--doc",
        str(replace_output_path),
        "--item",
        "para:2",
        "--text",
        "CLI appended acceptance text.",
    )
    append_output_path = _parse_output_path(append_result.stdout)
    write_result = _run_cli(
        golden_config_path,
        "write-cell",
        "--doc",
        str(golden_xlsx),
        "--sheet",
        "Budget2026",
        "--cell",
        "B2",
        "--value",
        "125001",
    )
    write_output_path = _parse_output_path(write_result.stdout)
    append_cell_result = _run_cli(
        golden_config_path,
        "append",
        "--doc",
        str(write_output_path),
        "--item",
        "sheet:Notes 2026!B5",
        "--text",
        "Acceptance note.",
    )
    append_cell_output_path = _parse_output_path(append_cell_result.stdout)
    doctor_result = _run_cli(golden_config_path, "doctor")

    assert index_result.returncode == 0, index_result.stderr
    assert "files_indexed=3" in index_result.stdout
    assert reindex_result.returncode == 0, reindex_result.stderr
    assert "files_indexed=1" in reindex_result.stdout
    assert list_result.returncode == 0, list_result.stderr
    assert "type=docx" in list_result.stdout
    assert "items=11" in list_result.stdout
    assert show_doc_result.returncode == 0, show_doc_result.stderr
    assert "item_count: 11" in show_doc_result.stdout
    assert show_item_result.returncode == 0, show_item_result.stderr
    assert "content_text: Supplier shall deliver by Friday." in show_item_result.stdout
    assert search_result.returncode == 0, search_result.stderr
    assert "para:3" in search_result.stdout
    assert locate_result.returncode == 0, locate_result.stderr
    assert "para:3" in locate_result.stdout
    assert read_result.returncode == 0, read_result.stderr
    assert "Supplier shall deliver by Friday." in read_result.stdout
    assert replace_result.returncode == 0, replace_result.stderr
    assert append_result.returncode == 0, append_result.stderr
    assert write_result.returncode == 0, write_result.stderr
    assert append_cell_result.returncode == 0, append_cell_result.stderr
    assert doctor_result.returncode == 0, doctor_result.stderr
    assert "Doctor Report" in doctor_result.stdout

    final_document = Document(str(append_output_path))
    assert final_document.paragraphs[1].text == "CLI acceptance text."
    assert final_document.paragraphs[2].text == "CLI appended acceptance text."

    workbook = load_workbook(append_cell_output_path)
    assert workbook["Budget2026"]["B2"].value == 125001
    assert workbook["Notes 2026"]["B5"].value == "Acceptance note."


def test_cli_acceptance_json_quiet_and_exit_codes(
    golden_config_path,
    golden_docx,
    golden_xlsx,
    golden_workspace,
    tmp_path,
) -> None:
    docs_dir = golden_workspace / "docs"
    _run_cli(golden_config_path, "index", str(docs_dir))

    quiet_result = _run_cli(golden_config_path, "search", "Supplier shall", "--quiet")
    list_json = _run_cli(golden_config_path, "list", "--json")
    show_json = _run_cli(golden_config_path, "show", "--doc", str(golden_docx), "--json")
    item_json = _run_cli(
        golden_config_path,
        "show",
        "--doc",
        str(golden_docx),
        "--item",
        "para:3",
        "--json",
    )
    doctor_json = _run_cli(golden_config_path, "doctor", "--json")
    invalid_args = _run_cli(golden_config_path, "locate", "--doc", str(golden_docx), "--slide", "1")
    not_found = _run_cli(
        golden_config_path,
        "show",
        "--doc",
        str(golden_docx),
        "--item",
        "para:999",
    )
    not_editable = _run_cli(
        golden_config_path,
        "append",
        "--doc",
        str(golden_xlsx),
        "--item",
        "sheet:Budget2026!B2",
        "--text",
        "blocked",
    )

    policy_config = tmp_path / "policy-office-agent.toml"
    policy_config.write_text(
        f"""
[offagent]
index_path = "{(tmp_path / 'state' / 'index.sqlite3').as_posix()}"
document_roots = ["{docs_dir.as_posix()}"]
allowed_roots = ["{docs_dir.as_posix()}"]
output_directory = "{(tmp_path / 'edited').as_posix()}"
output_roots = ["{(tmp_path / 'elsewhere').as_posix()}"]
""".strip()
    )
    _run_cli(policy_config, "index", str(golden_docx))
    policy_refused = _run_cli(
        policy_config,
        "replace",
        "--doc",
        str(golden_docx),
        "--item",
        "para:1",
        "--text",
        "blocked",
    )
    conflicting_modes = _run_cli(golden_config_path, "list", "--json", "--quiet")

    list_payload = json.loads(list_json.stdout)
    show_payload = json.loads(show_json.stdout)
    item_payload = json.loads(item_json.stdout)
    doctor_payload = json.loads(doctor_json.stdout)

    assert quiet_result.returncode == 0, quiet_result.stderr
    assert quiet_result.stdout == ""
    assert list_json.returncode == 0, list_json.stderr
    assert len(list_payload["documents"]) == 3
    assert any(document["item_count"] == 11 for document in list_payload["documents"])
    assert show_json.returncode == 0, show_json.stderr
    assert show_payload["item_count"] == 11
    assert item_json.returncode == 0, item_json.stderr
    assert item_payload["content_text"] == "Supplier shall deliver by Friday."
    assert doctor_json.returncode == 0, doctor_json.stderr
    assert doctor_payload["ok"] is True
    assert invalid_args.returncode == 2
    assert not_found.returncode == 3
    assert not_editable.returncode == 4
    assert policy_refused.returncode == 5
    assert conflicting_modes.returncode == 2
