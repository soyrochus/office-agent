from __future__ import annotations

from docx import Document

from offagent.adapters import docx_adapter


def test_extracts_docx_paragraphs_with_metadata(sample_docx) -> None:
    items = docx_adapter.extract_document(sample_docx)

    assert [item.item_id for item in items] == ["para:0", "para:1", "para:2", "para:3"]
    assert items[0].metadata["style_name"] == "Heading 1"
    assert items[0].metadata["is_heading"] is True
    assert items[2].content_text == ""
    assert all("Table text to ignore" not in item.content_text for item in items)


def test_replace_preserves_first_run_formatting(sample_docx) -> None:
    docx_adapter.replace_paragraph(sample_docx, "para:1", "Updated alpha paragraph.")

    document = Document(str(sample_docx))
    paragraph = document.paragraphs[1]
    assert paragraph.text == "Updated alpha paragraph."
    assert paragraph.runs[0].bold is True


def test_append_updates_empty_paragraph(sample_docx) -> None:
    docx_adapter.append_paragraph(sample_docx, "para:2", "Filled later.")

    document = Document(str(sample_docx))
    assert document.paragraphs[2].text == "Filled later."


def test_build_embedding_text_returns_paragraph_content(sample_docx) -> None:
    item = docx_adapter.extract_document(sample_docx)[3]

    assert docx_adapter.build_embedding_text(item, sample_docx) == item.content_text
