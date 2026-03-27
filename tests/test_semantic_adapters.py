from __future__ import annotations

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from offagent.adapters import docx_adapter, pptx_adapter, xlsx_adapter
from offagent.errors import TargetNotEditableError


def test_docx_semantic_helpers_preserve_block_order(sample_docx) -> None:
    blocks = docx_adapter.get_blocks(sample_docx)
    paragraphs = docx_adapter.get_paragraphs(sample_docx)
    tables = docx_adapter.get_tables(sample_docx)
    paragraph_bundle = docx_adapter.get_block_bundle(sample_docx, 0)
    table_bundle = docx_adapter.get_block_bundle(sample_docx, 4)

    assert [block.block_type for block in blocks] == [
        "paragraph",
        "paragraph",
        "paragraph",
        "paragraph",
        "table",
    ]
    assert paragraphs[0].style_name == "Heading 1"
    assert paragraphs[0].is_heading is True
    assert tables[0].rows == (("Table text to ignore",),)
    assert paragraph_bundle.paragraph is not None
    assert paragraph_bundle.paragraph.text == "Project Heading"
    assert table_bundle.table is not None
    assert table_bundle.table.rows[0][0] == "Table text to ignore"


def test_docx_semantic_writes_append_and_reject_table_replacement(sample_docx) -> None:
    output_path, block_index = docx_adapter.append_paragraph_block(
        sample_docx,
        "Tail semantic paragraph.",
        style_name="Heading 2",
    )

    document = Document(str(output_path))
    assert block_index == 5
    assert document.paragraphs[-1].text == "Tail semantic paragraph."
    assert document.paragraphs[-1].style.name == "Heading 2"

    with pytest.raises(TargetNotEditableError, match="table block replacement"):
        docx_adapter.replace_block(sample_docx, 4, "blocked")


def test_pptx_semantic_helpers_expose_slide_bundles_and_notes(sample_pptx) -> None:
    structure = pptx_adapter.get_presentation_structure(sample_pptx)
    bundle = pptx_adapter.get_slide_bundle(sample_pptx, 1)
    notes_text = pptx_adapter.get_slide_notes(sample_pptx, 2)

    assert [slide.slide_number for slide in structure] == [1, 2]
    assert bundle.notes_text == "Speaker notes: confirm launch owner."
    assert [block.shape_name for block in bundle.text_blocks[:3]] == [
        "Title 1",
        "TextBox 2",
        "TextBox 3",
    ]
    assert bundle.text_blocks[1].text == "Alpha launch overview\nSupplier shall present the rollout plan."
    assert notes_text == "Speaker notes: review action list."


def test_xlsx_semantic_helpers_snapshot_append_and_write_table(sample_xlsx, tmp_path) -> None:
    structure = xlsx_adapter.get_workbook_structure(sample_xlsx)
    snapshot = xlsx_adapter.get_sheet_snapshot(
        sample_xlsx,
        "Notes 2026",
        start_cell="A1",
        row_count=2,
        column_count=2,
    )
    append_output, row_number, coordinates = xlsx_adapter.append_row(
        sample_xlsx,
        "Notes 2026",
        values=["New note", "Finance"],
    )

    workbook = load_workbook(append_output)
    assert [sheet.sheet_name for sheet in structure.sheets] == ["Budget2026", "Notes 2026"]
    assert [cell.coordinate for cell in snapshot.cells] == ["A1", "B1", "A2", "B2"]
    assert snapshot.cells[0].display_value == "Supplier shall review variance."
    assert row_number == 3
    assert coordinates == ("A3", "B3")
    assert workbook["Notes 2026"]["A3"].value == "New note"
    assert workbook["Notes 2026"]["B3"].value == "Finance"

    mapped_workbook_path = tmp_path / "mapped.xlsx"
    mapped_workbook = Workbook()
    worksheet = mapped_workbook.active
    worksheet.title = "Tasks"
    worksheet["A1"] = "Task"
    worksheet["B1"] = "Owner"
    mapped_workbook.save(mapped_workbook_path)

    table_output, start_row, end_row = xlsx_adapter.write_table(
        mapped_workbook_path,
        "Tasks",
        records=[
            {"task": "Review budget", "owner": "Finance"},
            {"task": "Confirm launch", "owner": "Operations"},
        ],
        column_mapping={"task": "Task", "owner": "Owner"},
    )

    written_workbook = load_workbook(table_output)
    assert (start_row, end_row) == (2, 3)
    assert written_workbook["Tasks"]["A2"].value == "Review budget"
    assert written_workbook["Tasks"]["B3"].value == "Operations"
