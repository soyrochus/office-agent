from __future__ import annotations

import asyncio
import sys
from pathlib import Path

from docx import Document
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import stdio_client
from openpyxl import load_workbook


def _server_params(config_path: Path) -> StdioServerParameters:
    return StdioServerParameters(
        command=sys.executable,
        args=["-m", "offagent", "mcp", "--config", str(config_path)],
    )


def _run_mcp(config_path: Path, callback):
    async def runner():
        async with stdio_client(_server_params(config_path)) as (read_stream, write_stream):
            async with ClientSession(read_stream, write_stream) as session:
                await session.initialize()
                return await callback(session)

    return asyncio.run(runner())


async def _call_tool(session: ClientSession, name: str, arguments: dict | None = None):
    result = await session.call_tool(name, arguments=arguments)
    assert not result.isError, _tool_error_text(result)
    return result.structuredContent


def _tool_error_text(result) -> str:
    parts: list[str] = []
    for content in result.content:
        text = getattr(content, "text", None)
        if text is not None:
            parts.append(text)
    return "\n".join(parts)


def test_semantic_mcp_tools_are_discoverable_and_schema_backed(
    sample_docx,
    sample_pptx,
    sample_xlsx,
    config_path,
) -> None:
    def scenario(session: ClientSession):
        async def run():
            tools = await session.list_tools()
            tool_map = {tool.name: tool for tool in tools.tools}
            expected_tools = {
                "index_documents",
                "refresh_document",
                "list_documents",
                "search_objects",
                "search_documents",
                "get_structure",
                "get_section",
                "get_object",
                "list_children",
                "get_node",
                "write_node",
                "create_object",
                "update_object",
                "move_object",
                "copy_object",
                "batch_edit",
                "delete_object",
                "insert_content",
                "create_document",
                "add_content_block",
                "style_inline",
                "style_block",
                "set_structural_role",
                "xlsx_insert_rows",
                "xlsx_write_range",
                "xlsx_insert_columns",
                "xlsx_set_formula",
                "xlsx_merge_cells",
                "docx_get_tables",
                "docx_set_paragraph_style",
                "docx_insert_page_break",
                "docx_add_table",
                "docx_merge_table_cells",
                "pptx_add_slide",
                "pptx_duplicate_slide",
                "pptx_set_slide_layout",
                "pptx_add_text_shape",
            }

            assert expected_tools == set(tool_map)
            assert all(tool_map[name].inputSchema is not None for name in expected_tools)
            assert all(tool_map[name].outputSchema is not None for name in expected_tools)

            await _call_tool(
                session,
                "index_documents",
                {"paths": [str(sample_docx), str(sample_pptx), str(sample_xlsx)]},
            )
            documents = await _call_tool(session, "list_documents")
            docs_by_type = {document["file_type"]: document for document in documents["documents"]}

            docx_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_type["docx"]["document_id"]},
            )
            pptx_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_type["pptx"]["document_id"]},
            )
            xlsx_structure = await _call_tool(
                session,
                "get_structure",
                {"document_id": docs_by_type["xlsx"]["document_id"]},
            )

            table_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docs_by_type["docx"]["document_id"],
                    "section_id": docx_structure["sections"][-1]["locator"],
                },
            )
            slide_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docs_by_type["pptx"]["document_id"],
                    "section_id": pptx_structure["sections"][0]["locator"],
                },
            )
            sheet_section = await _call_tool(
                session,
                "get_section",
                {
                    "document_id": docs_by_type["xlsx"]["document_id"],
                    "section_id": xlsx_structure["sections"][1]["locator"],
                    "cell_range": "A1:B2",
                },
            )
            tables_result = await _call_tool(
                session,
                "docx_get_tables",
                {"document_id": docs_by_type["docx"]["document_id"]},
            )
            node_result = await _call_tool(
                session,
                "get_node",
                {"document_id": docs_by_type["docx"]["document_id"], "node_id": "para:0"},
            )
            write_result = await _call_tool(
                session,
                "write_node",
                {
                    "document_id": docs_by_type["docx"]["document_id"],
                    "node_id": "para:1",
                    "content": "Semantic MCP replacement.",
                },
            )
            insert_result = await _call_tool(
                session,
                "insert_content",
                {
                    "document_id": write_result["document_id"],
                    "content": "Semantic MCP paragraph.",
                    "style_name": "Heading 2",
                    "after_node_id": "para:1",
                },
            )
            append_rows = await _call_tool(
                session,
                "xlsx_insert_rows",
                {
                    "document_id": docs_by_type["xlsx"]["document_id"],
                    "sheet_name": "Notes 2026",
                    "rows": [["Semantic note", "Finance"]],
                },
            )

            return {
                "docx_structure": docx_structure,
                "pptx_structure": pptx_structure,
                "xlsx_structure": xlsx_structure,
                "table_section": table_section,
                "slide_section": slide_section,
                "sheet_section": sheet_section,
                "tables_result": tables_result,
                "node_result": node_result,
                "insert_result": insert_result,
                "append_rows": append_rows,
            }

        return run()

    result = _run_mcp(config_path, scenario)

    assert [section["section_type"] for section in result["docx_structure"]["sections"]] == [
        "paragraph",
        "paragraph",
        "paragraph",
        "paragraph",
        "table",
    ]
    assert [section["slide_number"] for section in result["pptx_structure"]["sections"]] == [1, 2]
    assert [section["sheet_name"] for section in result["xlsx_structure"]["sections"]] == [
        "Budget2026",
        "Notes 2026",
    ]
    assert result["table_section"]["rows"] == [["Table text to ignore"]]
    assert result["slide_section"]["notes_text"] == "Speaker notes: confirm launch owner."
    assert [cell["coordinate"] for cell in result["sheet_section"]["cells"]] == ["A1", "B1", "A2", "B2"]
    assert result["tables_result"]["tables"][0]["locator"] == "table:0:cell:0:0"
    assert result["node_result"]["text"] == "Project Heading"
    assert result["append_rows"]["first_row_locator"] == "sheet:Notes 2026!A3"

    inserted_document = Document(result["insert_result"]["output_path"])
    appended_workbook = load_workbook(result["append_rows"]["output_path"])
    assert inserted_document.paragraphs[2].text == "Semantic MCP paragraph."
    assert inserted_document.paragraphs[2].style.name == "Heading 2"
    assert appended_workbook["Notes 2026"]["A3"].value == "Semantic note"
    assert appended_workbook["Notes 2026"]["B3"].value == "Finance"


def test_semantic_mcp_reports_format_specific_errors(sample_docx, sample_pptx, sample_xlsx, config_path) -> None:
    def scenario(session: ClientSession):
        async def run():
            await _call_tool(
                session,
                "index_documents",
                {"paths": [str(sample_docx), str(sample_pptx), str(sample_xlsx)]},
            )
            documents = await _call_tool(session, "list_documents")
            docs_by_type = {document["file_type"]: document for document in documents["documents"]}

            unsupported_rows = await session.call_tool(
                "xlsx_insert_rows",
                {
                    "document_id": docs_by_type["docx"]["document_id"],
                    "sheet_name": "Sheet1",
                    "rows": [["blocked"]],
                },
            )
            assert unsupported_rows.isError
            assert "requires a .xlsx document" in _tool_error_text(unsupported_rows)

            unsupported_insert = await session.call_tool(
                "insert_content",
                {"document_id": docs_by_type["pptx"]["document_id"], "content": "blocked"},
            )
            assert unsupported_insert.isError
            assert "requires a .docx document" in _tool_error_text(unsupported_insert)

            unsupported_tables = await session.call_tool(
                "docx_get_tables",
                {"document_id": docs_by_type["xlsx"]["document_id"]},
            )
            assert unsupported_tables.isError
            assert "requires a .docx document" in _tool_error_text(unsupported_tables)

        return run()

    _run_mcp(config_path, scenario)
